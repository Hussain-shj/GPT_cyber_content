import base64
import hashlib
import hmac
import html
import io
import json
import math
import os
import re
import secrets
import shutil
import subprocess
import uuid
import threading
import tempfile
import time
import wave
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Literal, Any
from urllib.parse import unquote, urlparse, urlencode, quote

import psycopg
import httpx
from psycopg.rows import dict_row
from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response, RedirectResponse
from cryptography.fernet import Fernet, InvalidToken
from pydantic import BaseModel, Field
from openai import OpenAI

app = FastAPI(title="GPT Cyber Content API", version="0.52.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=False, allow_methods=["*"], allow_headers=["*"])
BASE_DIR = Path(__file__).resolve().parent
INDEX_FILE = BASE_DIR / "index.html"
MOBILE_JS = BASE_DIR / "mobile-download.js"
NEWS_SEARCH_JS = BASE_DIR / "news-search.js"
VISUAL_ALERT_JS = BASE_DIR / "visual-alert.js"
LINKEDIN_JS = BASE_DIR / "linkedin.js"
CYBER_SOURCES_FILE = BASE_DIR / "cyber_sources.json"
LINKEDIN_REFERENCE_PLAN_FILE = BASE_DIR / "linkedin_reference_plan.json"
AUTH_EXEMPT_PATHS = {"/health", "/auth/linkedin/callback"}

class ContentRequest(BaseModel):
    topic: str = Field(min_length=3, max_length=3000)
    domain: Literal["GRC", "Cybersecurity", "AI Governance", "Privacy"] = "GRC"
    post_type: Literal["Carousel", "Infographic", "Single Post"] = "Carousel"
    platform: Literal["Instagram", "LinkedIn", "Both"] = "Both"
    audience: str = "Government and enterprise cybersecurity professionals"
    language: Literal["Arabic", "English"] = "Arabic"
    slides: int = Field(default=6, ge=1, le=10)
    tone: str = "Professional, practical, executive-friendly"
    use_web_search: bool = False
    reference_week: int | None = Field(default=None, ge=1, le=12)
    reference_slot: int | None = Field(default=None, ge=1, le=4)

class ImageRequest(BaseModel):
    title: str
    body: str = ""
    slide_number: int = 1
    post_type: Literal["Carousel", "Infographic", "Single Post"] = "Single Post"
    domain: str = "GRC"
    visual_style: Literal["GRC Professional", "Cyber Pulse", "Executive Minimal", "Infographic"] = "GRC Professional"
    visual_direction: str = ""
    variant_index: int = Field(default=1, ge=1, le=3)

class NewsParseRequest(BaseModel):
    title: str = Field(min_length=3, max_length=500)
    news: str = Field(min_length=10, max_length=12000)

class NewsSearchRequest(BaseModel):
    days: int = Field(default=7, ge=1, le=30)
    limit: int = Field(default=8, ge=1, le=12)

class NewsVideoRequest(BaseModel):
    headline: str = Field(min_length=3, max_length=500)
    summary: str = Field(min_length=10, max_length=3000)
    threat_type: str = Field(default="خبر سيبراني", max_length=200)
    visual_brief: str = Field(default="", max_length=2000)
    style: Literal["Breaking News", "Cyber Awareness", "GRC"] = "Breaking News"
    duration: Literal[5, 10, 15] = 5

class VisualAlertRequest(BaseModel):
    title: str = Field(min_length=3, max_length=500)
    content: str = Field(min_length=10, max_length=12000)
    required_action: str = Field(min_length=3, max_length=4000)
    visual_style: Literal["Auto", "Cinematic AI", "SOC Operations", "Executive GRC", "Cyber Awareness"] = "Cinematic AI"
    threat_visual_style: Literal["Auto by Content", "Warning Screens", "Mobile Alerts", "System Errors and Updates", "Concerned User", "Anonymous Hacker", "Mixed Cyber Threats"] = "Auto by Content"
    video_count: Literal[0, 1, 3] = 1

class VisualIdeaRequest(BaseModel):
    idea: str = Field(min_length=3, max_length=6000)

class VisualApprovalRequest(BaseModel):
    asset_order: list[str] = Field(default_factory=list, max_length=6)
    motion_overlays: bool = False

class VisualSaveRequest(BaseModel):
    title: str = Field(min_length=3, max_length=500)
    content: str = Field(min_length=10, max_length=12000)
    required_action: str = Field(min_length=3, max_length=4000)
    formatted_content: str = Field(default="", max_length=16000)

class ArchivePost(BaseModel):
    id: str | None = None
    topic_id: int | None = None
    topic: str
    domain: str = "GRC"
    post_type: str = "Carousel"
    platform: str = "Both"
    content: dict[str, Any]

class LinkedInImage(BaseModel):
    image_b64: str = Field(min_length=1, max_length=18_000_000)
    image_mime_type: Literal["image/jpeg", "image/png"] = "image/jpeg"
    alt_text: str = Field(default="", max_length=120)

class LinkedInPublishRequest(BaseModel):
    text: str = Field(min_length=1, max_length=3000)
    image_b64: str = Field(default="", max_length=18_000_000)
    image_mime_type: Literal["image/jpeg", "image/png"] = "image/jpeg"
    images: list[LinkedInImage] = Field(default_factory=list, max_length=20)
    studio_post_id: str = Field(default="", max_length=200)

class LinkedInStudioPost(BaseModel):
    id: str = Field(min_length=1, max_length=200)
    week: int = Field(ge=1, le=12)
    day: str = Field(default="", max_length=100)
    pillar: str = Field(default="", max_length=200)
    objective: str = Field(default="", max_length=100)
    post_type: str = Field(default="Single Post", max_length=100)
    title: str = Field(min_length=1, max_length=800)
    status: str = Field(default="REVIEW", max_length=50)
    text: str = Field(default="", max_length=20_000)
    content: dict[str, Any] = Field(default_factory=dict)
    quality: dict[str, Any] = Field(default_factory=dict)

class LinkedInStudioAsset(BaseModel):
    id: str = Field(min_length=1, max_length=260)
    kind: Literal["image", "supporting_file"]
    name: str = Field(default="", max_length=500)
    mime_type: str = Field(default="application/octet-stream", max_length=200)
    data_b64: str = Field(min_length=1, max_length=28_000_000)
    sort_order: int = Field(default=0, ge=0, le=100)
    alt_text: str = Field(default="", max_length=300)

def database_url(): return os.getenv("DATABASE_URL")

VIDEO_JOBS: dict[str, dict[str, Any]] = {}
VIDEO_JOBS_LOCK = threading.Lock()
VISUAL_ALERT_JOBS: dict[str, dict[str, Any]] = {}
VISUAL_ALERT_JOBS_LOCK = threading.Lock()
def db_conn():
    if not database_url(): raise HTTPException(503, "DATABASE_URL is not configured")
    return psycopg.connect(database_url(), row_factory=dict_row)

def load_cyber_sources():
    try:
        return json.loads(CYBER_SOURCES_FILE.read_text(encoding="utf-8"))
    except Exception:
        return []

def load_linkedin_reference_plan():
    try:
        plan = json.loads(LINKEDIN_REFERENCE_PLAN_FILE.read_text(encoding="utf-8"))
        weeks = plan.get("weeks", [])
        if len(weeks) != 12 or any(len(week) != 4 for week in weeks):
            raise ValueError("Reference plan must contain 12 weeks with 4 topics each")
        return plan
    except Exception as exc:
        print("LinkedIn reference plan warning:", exc)
        return {"version": 0, "weeks": []}

def linkedin_reference_for(week: int | None, slot: int | None):
    if week is None or slot is None:
        return None
    plan = load_linkedin_reference_plan()
    try:
        item = plan["weeks"][week - 1][slot - 1]
        return item, plan.get("drive_folder_url", "")
    except (IndexError, KeyError, TypeError):
        raise HTTPException(422, "مرجع موضوع LinkedIn غير صالح")

def source_hosts():
    hosts = set()
    for src in load_cyber_sources():
        host = urlparse(src.get("url", "")).hostname
        if host:
            hosts.add(host.lower().removeprefix("www."))
    return hosts

def url_is_approved(url: str):
    try:
        host = (urlparse(url).hostname or "").lower().removeprefix("www.")
        return any(host == allowed or host.endswith("." + allowed) for allowed in source_hosts())
    except Exception:
        return False

def hash_password(password, salt=None):
    salt = salt or secrets.token_bytes(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt, 310000)
    return base64.b64encode(salt).decode(), base64.b64encode(digest).decode()

def verify_password(password, salt_b64, hash_b64):
    try:
        return hmac.compare_digest(hashlib.pbkdf2_hmac("sha256", password.encode(), base64.b64decode(salt_b64), 310000), base64.b64decode(hash_b64))
    except Exception:
        return False

def init_db():
    if not database_url(): return
    with psycopg.connect(database_url()) as conn:
        conn.execute("CREATE TABLE IF NOT EXISTS posts(id TEXT PRIMARY KEY,topic_id INTEGER,topic TEXT NOT NULL,domain TEXT NOT NULL,post_type TEXT NOT NULL,platform TEXT NOT NULL,content JSONB NOT NULL,created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW())")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_posts_topic_id ON posts(topic_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_posts_created_at ON posts(created_at DESC)")
        conn.execute("CREATE TABLE IF NOT EXISTS users(id TEXT PRIMARY KEY,username TEXT UNIQUE NOT NULL,password_salt TEXT NOT NULL,password_hash TEXT NOT NULL,is_active BOOLEAN NOT NULL DEFAULT TRUE,created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW())")
        conn.execute("CREATE TABLE IF NOT EXISTS visual_alert_archives(id TEXT PRIMARY KEY,title TEXT NOT NULL,content TEXT NOT NULL,required_action TEXT NOT NULL,formatted_content TEXT NOT NULL DEFAULT '',drive_video_id TEXT,drive_video_url TEXT,drive_text_id TEXT,drive_text_url TEXT,created_at TIMESTAMPTZ NOT NULL DEFAULT NOW())")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_visual_alert_archives_created_at ON visual_alert_archives(created_at DESC)")
        conn.execute("CREATE TABLE IF NOT EXISTS linkedin_connections(id TEXT PRIMARY KEY DEFAULT 'personal',member_id TEXT NOT NULL,member_name TEXT NOT NULL DEFAULT '',access_token_encrypted TEXT NOT NULL,expires_at TIMESTAMPTZ,connected_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW())")
        conn.execute("CREATE TABLE IF NOT EXISTS linkedin_publications(id TEXT PRIMARY KEY,post_urn TEXT NOT NULL,post_url TEXT NOT NULL,has_image BOOLEAN NOT NULL DEFAULT FALSE,created_at TIMESTAMPTZ NOT NULL DEFAULT NOW())")
        conn.execute("CREATE TABLE IF NOT EXISTS linkedin_studio_posts(id TEXT PRIMARY KEY,week INTEGER NOT NULL CHECK(week BETWEEN 1 AND 12),day TEXT NOT NULL DEFAULT '',pillar TEXT NOT NULL DEFAULT '',objective TEXT NOT NULL DEFAULT '',post_type TEXT NOT NULL DEFAULT 'Single Post',title TEXT NOT NULL,status TEXT NOT NULL DEFAULT 'REVIEW',text TEXT NOT NULL DEFAULT '',content JSONB NOT NULL DEFAULT '{}'::jsonb,quality JSONB NOT NULL DEFAULT '{}'::jsonb,published_url TEXT NOT NULL DEFAULT '',published_at TIMESTAMPTZ,created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW())")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_linkedin_studio_posts_week ON linkedin_studio_posts(week,created_at)")
        conn.execute("CREATE TABLE IF NOT EXISTS linkedin_studio_assets(id TEXT PRIMARY KEY,post_id TEXT NOT NULL REFERENCES linkedin_studio_posts(id) ON DELETE CASCADE,kind TEXT NOT NULL CHECK(kind IN ('image','supporting_file')),name TEXT NOT NULL DEFAULT '',mime_type TEXT NOT NULL,data BYTEA NOT NULL,sort_order INTEGER NOT NULL DEFAULT 0,alt_text TEXT NOT NULL DEFAULT '',created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW())")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_linkedin_studio_assets_post ON linkedin_studio_assets(post_id,kind,sort_order)")
        conn.commit()

def bootstrap_user():
    if not database_url(): return
    username = os.getenv("AUTH_BOOTSTRAP_USERNAME", "").strip()
    password = os.getenv("AUTH_BOOTSTRAP_PASSWORD", "")
    if not username or not password: return
    with psycopg.connect(database_url(), row_factory=dict_row) as conn:
        if conn.execute("SELECT COUNT(*) AS total FROM users").fetchone()["total"] == 0:
            salt, pwh = hash_password(password)
            conn.execute("INSERT INTO users(id,username,password_salt,password_hash,is_active) VALUES(%s,%s,%s,%s,TRUE)", (str(uuid.uuid4()), username, salt, pwh))
            conn.commit()

def user_count():
    if not database_url(): return 0
    try:
        with psycopg.connect(database_url(), row_factory=dict_row) as conn:
            return conn.execute("SELECT COUNT(*) AS total FROM users WHERE is_active=TRUE").fetchone()["total"]
    except Exception:
        return 0

def authenticate_basic(header):
    if not header or not header.startswith("Basic "): return False
    try:
        username, password = base64.b64decode(header.split(" ", 1)[1]).decode().split(":", 1)
        with psycopg.connect(database_url(), row_factory=dict_row) as conn:
            u = conn.execute("SELECT * FROM users WHERE username=%s", (username,)).fetchone()
        return bool(u and u["is_active"] and verify_password(password, u["password_salt"], u["password_hash"]))
    except Exception:
        return False

@app.middleware("http")
async def auth(request: Request, call_next):
    if request.url.path in AUTH_EXEMPT_PATHS: return await call_next(request)
    if not database_url(): return JSONResponse({"detail":"DATABASE_URL is not configured"}, status_code=503)
    if user_count() == 0: return JSONResponse({"detail":"No active users found. Set AUTH_BOOTSTRAP_USERNAME and AUTH_BOOTSTRAP_PASSWORD in Railway Variables, then redeploy."}, status_code=503)
    if not authenticate_basic(request.headers.get("Authorization")):
        return Response(status_code=401, headers={"WWW-Authenticate":"Basic realm=\"GPT Cyber Content\", charset=\"UTF-8\""})
    return await call_next(request)

@app.on_event("startup")
def startup():
    try: init_db(); bootstrap_user()
    except Exception as e: print("Database startup warning:", e)

def extract_json(text):
    cleaned = re.sub(r"^```json\s*|^```\s*|\s*```$", "", text.strip(), flags=re.I|re.S)
    return json.loads(cleaned)

@app.get("/", include_in_schema=False)
def web_app(): return FileResponse(INDEX_FILE, media_type="text/html")

@app.get("/mobile-download.js", include_in_schema=False)
def mobile_js():
    base = MOBILE_JS.read_text(encoding="utf-8") if MOBILE_JS.exists() else ""
    search = NEWS_SEARCH_JS.read_text(encoding="utf-8") if NEWS_SEARCH_JS.exists() else ""
    visual_alert = VISUAL_ALERT_JS.read_text(encoding="utf-8") if VISUAL_ALERT_JS.exists() else ""
    linkedin = LINKEDIN_JS.read_text(encoding="utf-8") if LINKEDIN_JS.exists() else ""
    return Response(content=base + "\n\n" + search + "\n\n" + visual_alert + "\n\n" + linkedin, media_type="application/javascript", headers={"Cache-Control":"no-store, max-age=0"})

@app.get("/health")
def health():
    return {
        "status":"ok", "version":"0.52.0", "openai_configured":bool(os.getenv("OPENAI_API_KEY")),
        "gemini_configured":bool(os.getenv("GEMINI_API_KEY")),
        "image_provider":"google_nano_banana_2_with_openai_fallback" if os.getenv("GEMINI_API_KEY") and os.getenv("OPENAI_API_KEY") else "google_nano_banana_2" if os.getenv("GEMINI_API_KEY") else "openai" if os.getenv("OPENAI_API_KEY") else "unconfigured",
        "image_model":os.getenv("GEMINI_IMAGE_MODEL", "gemini-3.1-flash-image"),
        "database_connected":bool(database_url()), "active_users":user_count(), "google_drive_configured":all(os.getenv(key) for key in ("GOOGLE_DRIVE_CLIENT_ID","GOOGLE_DRIVE_CLIENT_SECRET","GOOGLE_DRIVE_REFRESH_TOKEN")), "news_parser":"source-date-verified-v4",
        "news_artwork":"two-line-cve-layout-v12", "news_search":"approved-sources-v1", "news_sources":len(load_cyber_sources()),
        "bytez_video_configured":bool(os.getenv("BYTEZ_API_KEY")),
        "bytez_video_model":os.getenv("BYTEZ_VIDEO_MODEL", "automatic"),
        "visual_alert_editor":"fixed-brand-outro-v24", "brand_outro":True, "gemini_tts_model":os.getenv("GEMINI_TTS_MODEL", "gemini-3.1-flash-tts-preview"),
        "remotion_runtime_ready":bool(shutil.which("node") and (BASE_DIR / "node_modules" / "@remotion" / "renderer").exists()),
        "linkedin_reference_topics":sum(len(week) for week in load_linkedin_reference_plan().get("weeks", []))
    }

@app.get("/api/linkedin/reference-plan")
def linkedin_reference_plan():
    plan = load_linkedin_reference_plan()
    if not plan.get("weeks"):
        raise HTTPException(503, "مكتبة مراجع LinkedIn غير متاحة")
    return plan

def _linkedin_config():
    client_id = os.getenv("LINKEDIN_CLIENT_ID", "").strip()
    client_secret = os.getenv("LINKEDIN_CLIENT_SECRET", "").strip()
    redirect_uri = os.getenv("LINKEDIN_REDIRECT_URI", "").strip()
    if not all((client_id, client_secret, redirect_uri)):
        raise HTTPException(503, "أضف LINKEDIN_CLIENT_ID وLINKEDIN_CLIENT_SECRET وLINKEDIN_REDIRECT_URI في Railway")
    return client_id, client_secret, redirect_uri

def _linkedin_fernet():
    _, secret, _ = _linkedin_config()
    key = base64.urlsafe_b64encode(hashlib.sha256(("cyberpulse-linkedin:" + secret).encode()).digest())
    return Fernet(key)

def _linkedin_token():
    with db_conn() as c:
        row = c.execute("SELECT * FROM linkedin_connections WHERE id='personal'").fetchone()
    if not row: raise HTTPException(409, "حساب LinkedIn غير متصل")
    if row.get("expires_at") and row["expires_at"] <= datetime.now(timezone.utc):
        raise HTTPException(401, "انتهت صلاحية ربط LinkedIn. أعد الاتصال بالحساب")
    try: return _linkedin_fernet().decrypt(row["access_token_encrypted"].encode()).decode(), row
    except InvalidToken: raise HTTPException(500, "تعذر قراءة اتصال LinkedIn. أعد ربط الحساب")

@app.get("/auth/linkedin/start")
def linkedin_start():
    client_id, secret, redirect_uri = _linkedin_config()
    issued = str(int(time.time()))
    nonce = secrets.token_urlsafe(18)
    payload = issued + "." + nonce
    signature = hmac.new(secret.encode(), payload.encode(), hashlib.sha256).hexdigest()
    state = base64.urlsafe_b64encode((payload + "." + signature).encode()).decode().rstrip("=")
    query = urlencode({"response_type":"code","client_id":client_id,"redirect_uri":redirect_uri,"state":state,"scope":"openid profile w_member_social"})
    return RedirectResponse("https://www.linkedin.com/oauth/v2/authorization?" + query)

@app.get("/auth/linkedin/callback")
def linkedin_callback(code: str = "", state: str = "", error: str = ""):
    if error: return RedirectResponse("/#linkedin?linkedin=denied")
    client_id, secret, redirect_uri = _linkedin_config()
    try:
        decoded = base64.urlsafe_b64decode(state + "=" * (-len(state) % 4)).decode()
        issued, nonce, signature = decoded.split(".", 2)
        expected = hmac.new(secret.encode(), f"{issued}.{nonce}".encode(), hashlib.sha256).hexdigest()
        if not hmac.compare_digest(signature, expected) or abs(int(time.time()) - int(issued)) > 900: raise ValueError()
    except Exception: raise HTTPException(400, "حالة OAuth غير صالحة أو منتهية")
    if not code: raise HTTPException(400, "لم يُرجع LinkedIn رمز التفويض")
    with httpx.Client(timeout=40) as client:
        token_response = client.post("https://www.linkedin.com/oauth/v2/accessToken", data={"grant_type":"authorization_code","code":code,"redirect_uri":redirect_uri,"client_id":client_id,"client_secret":secret})
        if token_response.status_code >= 400: raise HTTPException(502, f"تعذر الحصول على رمز LinkedIn: {token_response.text[:500]}")
        token_data = token_response.json(); token = token_data.get("access_token", "")
        profile_response = client.get("https://api.linkedin.com/v2/userinfo", headers={"Authorization":f"Bearer {token}"})
        if profile_response.status_code >= 400: raise HTTPException(502, f"تعذر قراءة حساب LinkedIn: {profile_response.text[:500]}")
        profile = profile_response.json()
    member_id = str(profile.get("sub", "")).strip()
    if not token or not member_id: raise HTTPException(502, "لم يُرجع LinkedIn بيانات الحساب كاملة")
    expires_at = datetime.fromtimestamp(time.time() + int(token_data.get("expires_in", 5184000)), timezone.utc)
    encrypted = _linkedin_fernet().encrypt(token.encode()).decode()
    name = str(profile.get("name") or "حساب LinkedIn")[:300]
    with db_conn() as c:
        c.execute("INSERT INTO linkedin_connections(id,member_id,member_name,access_token_encrypted,expires_at) VALUES('personal',%s,%s,%s,%s) ON CONFLICT(id) DO UPDATE SET member_id=EXCLUDED.member_id,member_name=EXCLUDED.member_name,access_token_encrypted=EXCLUDED.access_token_encrypted,expires_at=EXCLUDED.expires_at,updated_at=NOW()", (member_id,name,encrypted,expires_at)); c.commit()
    return RedirectResponse("/#linkedin?linkedin=connected")

@app.get("/api/linkedin/status")
def linkedin_status():
    configured = all(os.getenv(k) for k in ("LINKEDIN_CLIENT_ID","LINKEDIN_CLIENT_SECRET","LINKEDIN_REDIRECT_URI"))
    if not configured: return {"configured":False,"connected":False}
    with db_conn() as c: row = c.execute("SELECT member_name,expires_at,connected_at FROM linkedin_connections WHERE id='personal'").fetchone()
    connected = bool(row and (not row["expires_at"] or row["expires_at"] > datetime.now(timezone.utc)))
    return {"configured":True,"connected":connected,"member_name":row["member_name"] if row else "","expires_at":row["expires_at"] if row else None}

@app.post("/api/linkedin/disconnect")
def linkedin_disconnect():
    with db_conn() as c: c.execute("DELETE FROM linkedin_connections WHERE id='personal'"); c.commit()
    return {"disconnected":True}

@app.get("/api/linkedin/studio/posts")
def linkedin_studio_posts(week: int | None = None):
    query = "SELECT id,week,day,pillar,objective,post_type AS type,title,status,text,content,quality,published_url,published_at,created_at,updated_at FROM linkedin_studio_posts"
    params: tuple[Any, ...] = ()
    if week is not None:
        if week < 1 or week > 12: raise HTTPException(400, "رقم الأسبوع يجب أن يكون من 1 إلى 12")
        query += " WHERE week=%s"; params = (week,)
    query += " ORDER BY week,created_at"
    with db_conn() as c: return c.execute(query, params).fetchall()

@app.post("/api/linkedin/studio/posts")
def linkedin_studio_post_save(p: LinkedInStudioPost):
    with db_conn() as c:
        c.execute("""INSERT INTO linkedin_studio_posts(id,week,day,pillar,objective,post_type,title,status,text,content,quality)
        VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s::jsonb,%s::jsonb)
        ON CONFLICT(id) DO UPDATE SET week=EXCLUDED.week,day=EXCLUDED.day,pillar=EXCLUDED.pillar,objective=EXCLUDED.objective,post_type=EXCLUDED.post_type,title=EXCLUDED.title,status=EXCLUDED.status,text=EXCLUDED.text,content=EXCLUDED.content,quality=EXCLUDED.quality,updated_at=NOW()""",
        (p.id,p.week,p.day,p.pillar,p.objective,p.post_type,p.title,p.status,p.text,json.dumps(p.content,ensure_ascii=False),json.dumps(p.quality,ensure_ascii=False))); c.commit()
    return {"saved":True,"id":p.id}

@app.delete("/api/linkedin/studio/posts/{post_id}")
def linkedin_studio_post_delete(post_id: str):
    with db_conn() as c:
        cur = c.execute("DELETE FROM linkedin_studio_posts WHERE id=%s", (post_id,)); c.commit()
    return {"deleted":cur.rowcount>0}

@app.get("/api/linkedin/studio/posts/{post_id}/assets")
def linkedin_studio_assets(post_id: str):
    with db_conn() as c:
        rows = c.execute("SELECT id,kind,name,mime_type,data,sort_order,alt_text,created_at,updated_at FROM linkedin_studio_assets WHERE post_id=%s ORDER BY kind,sort_order,created_at", (post_id,)).fetchall()
    return [{**{k:v for k,v in row.items() if k != "data"},"data_b64":base64.b64encode(bytes(row["data"])).decode()} for row in rows]

@app.post("/api/linkedin/studio/posts/{post_id}/assets")
def linkedin_studio_asset_save(post_id: str, asset: LinkedInStudioAsset):
    try: payload = base64.b64decode(asset.data_b64, validate=True)
    except Exception: raise HTTPException(400, "بيانات الملف غير صالحة")
    if len(payload) > 20 * 1024 * 1024: raise HTTPException(413, "حجم الملف يتجاوز 20 MB")
    with db_conn() as c:
        if not c.execute("SELECT 1 FROM linkedin_studio_posts WHERE id=%s", (post_id,)).fetchone(): raise HTTPException(404, "منشور LinkedIn غير موجود")
        c.execute("""INSERT INTO linkedin_studio_assets(id,post_id,kind,name,mime_type,data,sort_order,alt_text)
        VALUES(%s,%s,%s,%s,%s,%s,%s,%s)
        ON CONFLICT(id) DO UPDATE SET kind=EXCLUDED.kind,name=EXCLUDED.name,mime_type=EXCLUDED.mime_type,data=EXCLUDED.data,sort_order=EXCLUDED.sort_order,alt_text=EXCLUDED.alt_text,updated_at=NOW()""",
        (asset.id,post_id,asset.kind,asset.name,asset.mime_type,payload,asset.sort_order,asset.alt_text)); c.commit()
    return {"saved":True,"id":asset.id,"size":len(payload)}

@app.get("/api/linkedin/studio/posts/{post_id}/assets/{asset_id}/download")
def linkedin_studio_asset_download(post_id: str, asset_id: str):
    with db_conn() as c:
        row = c.execute("SELECT name,mime_type,data FROM linkedin_studio_assets WHERE id=%s AND post_id=%s", (asset_id,post_id)).fetchone()
    if not row: raise HTTPException(404, "الملف غير موجود")
    filename = row["name"] or "cyberpulse-file"
    return Response(content=bytes(row["data"]),media_type=row["mime_type"],headers={"Content-Disposition":f"attachment; filename*=UTF-8''{quote(filename)}"})

@app.delete("/api/linkedin/studio/posts/{post_id}/assets/{asset_id}")
def linkedin_studio_asset_delete(post_id: str, asset_id: str):
    with db_conn() as c:
        cur = c.execute("DELETE FROM linkedin_studio_assets WHERE id=%s AND post_id=%s", (asset_id,post_id)); c.commit()
    return {"deleted":cur.rowcount>0}

def _linkedin_rtl_text(text: str):
    normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    if not re.search(r"[\u0600-\u06ff\u0750-\u077f\u08a0-\u08ff]", normalized):
        return normalized
    return "\n".join(
        line if not line.strip() or line.startswith("\u200f") else "\u200f" + line
        for line in normalized.split("\n")
    )

@app.post("/api/linkedin/publish")
def linkedin_publish(req: LinkedInPublishRequest):
    post_text = _linkedin_rtl_text(req.text)
    if len(post_text) > 3000:
        raise HTTPException(400, f"النص يتجاوز حد LinkedIn بمقدار {len(post_text) - 3000} حرفًا بعد ضبط الاتجاه. اختصره قبل النشر.")
    token, connection = _linkedin_token()
    author = "urn:li:person:" + connection["member_id"]
    headers = {"Authorization":f"Bearer {token}","Linkedin-Version":os.getenv("LINKEDIN_API_VERSION", "202608"),"X-Restli-Protocol-Version":"2.0.0","Content-Type":"application/json"}
    content = None
    images = list(req.images)
    if not images and req.image_b64:
        images.append(LinkedInImage(image_b64=req.image_b64, image_mime_type=req.image_mime_type))
    uploaded_images = []
    with httpx.Client(timeout=120) as client:
        for index, image in enumerate(images, 1):
            try: image_bytes = base64.b64decode(image.image_b64, validate=True)
            except Exception: raise HTTPException(400, f"بيانات الصورة {index} غير صالحة")
            if len(image_bytes) > 10 * 1024 * 1024: raise HTTPException(413, f"حجم الصورة {index} يتجاوز 10 MB")
            init = client.post("https://api.linkedin.com/rest/images?action=initializeUpload", headers=headers, json={"initializeUploadRequest":{"owner":author}})
            if init.status_code >= 400: raise HTTPException(init.status_code, f"تعذر تهيئة صورة LinkedIn رقم {index}: {init.text[:700]}")
            value = init.json().get("value", {}); upload_url = value.get("uploadUrl"); image_urn = value.get("image")
            if not upload_url or not image_urn: raise HTTPException(502, f"لم يُرجع LinkedIn رابط رفع الصورة {index}")
            upload = client.put(upload_url, content=image_bytes, headers={"Content-Type":image.image_mime_type})
            if upload.status_code >= 400: raise HTTPException(upload.status_code, f"تعذر رفع صورة LinkedIn رقم {index}: {upload.text[:500]}")
            item = {"id":image_urn}
            if image.alt_text: item["altText"] = image.alt_text
            uploaded_images.append(item)
        if len(uploaded_images) == 1:
            content = {"media":{"title":"نبض سيبراني | CYBER PULSE","id":uploaded_images[0]["id"]}}
        elif len(uploaded_images) > 1:
            content = {"multiImage":{"images":uploaded_images}}
        body = {"author":author,"commentary":post_text,"visibility":"PUBLIC","distribution":{"feedDistribution":"MAIN_FEED","targetEntities":[],"thirdPartyDistributionChannels":[]},"lifecycleState":"PUBLISHED","isReshareDisabledByAuthor":False}
        if content: body["content"] = content
        response = client.post("https://api.linkedin.com/rest/posts", headers=headers, json=body)
    if response.status_code >= 400: raise HTTPException(response.status_code, f"فشل النشر على LinkedIn: {response.text[:900]}")
    post_urn = response.headers.get("x-restli-id", "")
    post_url = "https://www.linkedin.com/feed/update/" + post_urn + "/" if post_urn else "https://www.linkedin.com/feed/"
    with db_conn() as c:
        c.execute("INSERT INTO linkedin_publications(id,post_urn,post_url,has_image) VALUES(%s,%s,%s,%s)", (str(uuid.uuid4()),post_urn,post_url,bool(content))); c.commit()
        if req.studio_post_id:
            c.execute("UPDATE linkedin_studio_posts SET status='PUBLISHED',published_url=%s,published_at=NOW(),updated_at=NOW() WHERE id=%s", (post_url,req.studio_post_id)); c.commit()
    return {"published":True,"post_urn":post_urn,"post_url":post_url,"image_count":len(uploaded_images)}

def _visual_job_update(job_id: str, **values):
    with VISUAL_ALERT_JOBS_LOCK:
        if job_id in VISUAL_ALERT_JOBS: VISUAL_ALERT_JOBS[job_id].update(values)

@app.post("/api/visual-alert/analyze-idea")
def analyze_visual_alert_idea(req: VisualIdeaRequest):
    if not os.getenv("OPENAI_API_KEY"): raise HTTPException(400, "OPENAI_API_KEY is not configured")
    prompt = f'''You are the Arabic content strategist for a UAE cybersecurity awareness channel.
Turn the user's rough idea into editable input fields for a vertical awareness segment of 32-44 seconds; a fixed 10-second branded outro is appended later.
Write clear Modern Standard Arabic suitable for UAE government and enterprise audiences. Keep the tone practical, calm, concise and human-centered.
Do not invent CVEs, dates, vendors, incidents, statistics, laws, exploitation claims, product versions or technical facts that the user did not supply. If the idea is general awareness, develop it using durable best practices without pretending a specific incident occurred.
The content must be 70-105 Arabic words and explain the issue, why it matters, and the safe behavior. End the content with a short final paragraph beginning exactly with "الخلاصة:" that gives one memorable practical takeaway. The required_action must contain 3-5 short actionable lines separated by newlines, with no numbering or Markdown. The title must be compelling but factual, maximum 10 words.
Choose visual_style from exactly: Cinematic AI, Auto, SOC Operations, Executive GRC, Cyber Awareness. Prefer Cyber Awareness for personal behavior, families, phishing, passwords, social media or AI-use topics; Executive GRC for governance and management; SOC Operations only for genuinely technical operational incidents; otherwise Cinematic AI.
Return ONLY valid JSON with: title, content, required_action, visual_style.

USER IDEA:
{req.idea}'''
    try:
        client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
        data = extract_json(client.responses.create(model=os.getenv("OPENAI_MODEL", "gpt-5"), input=prompt, store=False).output_text)
        allowed_styles = {"Cinematic AI", "Auto", "SOC Operations", "Executive GRC", "Cyber Awareness"}
        title = " ".join(str(data.get("title", "")).split())[:500]
        content = str(data.get("content", "")).strip()[:12000]
        required_action = str(data.get("required_action", "")).strip()[:4000]
        if len(title) < 3 or len(content) < 10 or len(required_action) < 3: raise ValueError("لم تكتمل الحقول المقترحة")
        style = str(data.get("visual_style", "Cinematic AI"))
        return {"title":title, "content":content, "required_action":required_action, "visual_style":style if style in allowed_styles else "Cinematic AI"}
    except Exception as exc:
        raise HTTPException(500, f"تعذر تحليل الفكرة: {str(exc)[:700]}")

@app.post("/api/visual-alert/social-copy")
def create_visual_alert_social_copy(req: VisualAlertRequest):
    if not os.getenv("OPENAI_API_KEY"): raise HTTPException(400, "OPENAI_API_KEY is not configured")
    prompt = f'''You format Arabic cybersecurity content for LinkedIn and Instagram.
Use ONLY the supplied title, content and required actions. Never invent statistics, dates, standards, CVEs, vendors, incidents, claims, recommendations or facts.
Create one polished plain-text Arabic post ready to paste, following this exact visual rhythm:
1. First line: one relevant emoji followed by the factual title.
2. One or two short explanatory paragraphs with intentional line breaks.
3. A short section label appropriate to the content, followed by each supplied action on its own line. Use number-keycap emojis 1️⃣ 2️⃣ 3️⃣ 4️⃣ 5️⃣ in the original action order.
4. If the supplied content contains distinct supported focus areas or facts, add a short section label and 3-6 lines, each beginning with one relevant emoji such as 🔐 ⚠️ 🔍 🛡️ 📱 💻 📧 🔄. Omit this section if the source does not support it.
5. End with a concise conclusion. If the content already contains "الخلاصة:", preserve its meaning without duplicating it.
6. Add 8-12 concise relevant Arabic and English hashtags on the final lines. Always include #نبض_سيبراني and #الأمن_السيبراني.
Keep paragraphs short and mobile-friendly. Use no Markdown headings, asterisks, bullet characters, tables or fabricated citations. Preserve English technical terms where useful.
Return ONLY valid JSON: {{"social_copy":"..."}}

TITLE:
{req.title}

CONTENT:
{req.content}

REQUIRED ACTIONS:
{req.required_action}'''
    try:
        client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
        data = extract_json(client.responses.create(model=os.getenv("OPENAI_MODEL", "gpt-5"), input=prompt, store=False).output_text)
        social_copy = str(data.get("social_copy", "")).strip()
        if len(social_copy) < 20: raise ValueError("لم يُرجع OpenAI محتوى منسقًا")
        return {"social_copy":social_copy[:12000]}
    except Exception as exc:
        raise HTTPException(500, f"تعذر إنشاء محتوى مواقع التواصل: {str(exc)[:700]}")

def _visual_script(req: VisualAlertRequest) -> dict[str, Any]:
    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
    prompt = f'''You are a professional cybersecurity short-form video editor.
Transform ONLY the supplied Arabic cybersecurity alert into a concise vertical video script. Never invent CVEs, severity, versions, vendors, attack vectors, exploitation status, patches, affected systems, IOCs, or recommendations not explicitly supplied.
Target 32-44 seconds and never exceed 46 seconds. A fixed 10-second branded outro is appended after these scenes, and the complete exported video must remain under 58 seconds. The combined voiceText across ALL scenes must be no more than 82 Arabic words. Use the minimum number of scenes needed. Arabic RTL. On-screen text is maximum 9 words and 2 lines. Voice text is natural, concise, and not a verbatim copy of on-screen text.
Build a visually diverse sequence driven by the meaning of each scene. Do NOT default every scene to a SOC, control room, server room, or wall of screens. Choose a different appropriate real-world setting for each scene from possibilities such as a private employee office, open-plan workplace, executive office, meeting room, government service counter, reception, remote-work desk, home living room, family using phones or a laptop, café, airport or travel setting, data center, technical workshop, or SOC. Use a SOC or server room only when the supplied alert specifically supports it. A family or home scene is allowed only when the content concerns personal cyber safety, children, parents, home devices, scams, passwords, social media, or public awareness.
A scene may instead contain no people and use a realistic screen-capture-style or object-focused visual when that communicates the supplied content better. Appropriate examples include a red compromised computer screen, phishing email, suspicious link, fake sign-in page, account takeover warning, malware alert, unusual login activity, compromised cloud session, fraudulent message, unsafe attachment, security update, patch deployment, locked device, or incident dashboard. The screen must show the exact supported mechanism through recognizable interface structure, cursor focus, warning state, highlighted suspicious element, or security status—without inventing names, credentials, messages, CVEs, brands, or technical facts. Use screen capture only when directly relevant, and vary it with other scene types rather than making every scene a screen.
An anonymous hooded attacker at a laptop may be used as a dark navy/cyan cinematic silhouette similar to a serious cybersecurity editorial photograph, but ONLY when the supplied content directly discusses an attacker, active hacking, compromise, intrusion, malware operation or theft. Keep the face obscured, do not identify a nationality, organization or real person, and do not use this image as a generic decoration for ordinary awareness topics.
Every visualSuggestion must name one concrete setting or screen type, subjects if any, exact action, and relevant device; settings and compositions must not repeat across scenes. Keep people respectful, professional, realistic, culturally accurate, and relevant to the supplied alert.
Return ONLY valid JSON with videoTitle, estimatedDuration, and scenes. Each scene must contain id, type (intro/headline/content/risk/action/outro), duration (integer seconds), onScreenText, voiceText, subtitleEnglish, visualMode (environment/screen_capture/device_closeup/object_only/hacker_scene), visualSetting, visualSuggestion. subtitleEnglish must be a faithful, concise English translation of voiceText, suitable for a single subtitle line; preserve product names, CVEs, versions and technical meaning exactly. The required action must be communicated clearly near the end. The FINAL scene must always be type outro and provide a concise conclusion or memorable takeaway based only on the supplied content.

ALERT TITLE:
{req.title}

ALERT CONTENT:
{req.content}

REQUIRED ACTION:
{req.required_action}

SELECTED VISUAL STYLE:
{req.visual_style}

PREFERRED THREAT EDITORIAL STYLE:
{req.threat_visual_style}
Treat this as a preferred visual option for relevant scenes, not a requirement to repeat it in every asset. When "Mixed Cyber Threats" is selected, deliberately vary among red warning screens, mobile alerts, system errors or updates, a concerned device user, and an anonymous attacker only where factually appropriate.'''
    raw = client.responses.create(model=os.getenv("OPENAI_MODEL", "gpt-5"), input=prompt, store=False).output_text
    data = extract_json(raw)
    scenes = data.get("scenes") if isinstance(data, dict) else None
    if not isinstance(scenes, list) or not scenes: raise ValueError("لم يُرجع OpenAI مشاهد صالحة")
    clean = []
    for i, scene in enumerate(scenes[:8]):
        if not isinstance(scene, dict): continue
        clean.append({
            "id":f"scene-{i+1}", "type":str(scene.get("type", "content"))[:30],
            "duration":max(3, min(15, int(scene.get("duration", 6)))),
            "onScreenText":" ".join(str(scene.get("onScreenText", "")).split()[:9]),
            "voiceText":str(scene.get("voiceText", "")).strip()[:700],
            "subtitleEnglish":" ".join(str(scene.get("subtitleEnglish", "")).strip().split())[:500],
            "visualMode":str(scene.get("visualMode", "environment")).strip()[:40],
            "visualSetting":str(scene.get("visualSetting", "")).strip()[:160],
            "visualSuggestion":str(scene.get("visualSuggestion", "")).strip()[:500],
        })
    if not clean: raise ValueError("تعذر تكوين مشاهد الفيديو")
    remaining = 105
    for i, scene in enumerate(clean):
        words = scene["voiceText"].split()
        later = len(clean) - i - 1
        allowance = min(20, max(6, remaining - later * 12))
        scene["voiceText"] = " ".join(words[:allowance])
        remaining -= min(len(words), allowance)
    return {"videoTitle":str(data.get("videoTitle", req.title))[:300], "scenes":clean}

def _veo_prompt(scene: dict[str, Any], req: VisualAlertRequest) -> str:
    styles = {
        "Cinematic AI":"premium cinematic realism, dramatic controlled lighting, shallow depth of field, slow dolly camera, polished film color grade",
        "SOC Operations":"realistic UAE cybersecurity operations aesthetic when the scene actually requires it, focused response activity, cool blue practical lighting",
        "Executive GRC":"premium UAE government executive environment, risk and governance meeting, elegant architecture, restrained corporate lighting",
        "Cyber Awareness":"human-centered UAE workplace cybersecurity awareness scene, clear everyday action, warm professional lighting",
    }
    selected = req.visual_style
    if selected == "Auto":
        selected = "Executive GRC" if scene.get("type") == "action" else "Cyber Awareness" if scene.get("type") in {"risk","content"} else "Cinematic AI"
    threat_styles = {
        "Auto by Content":"Choose the most accurate threat visual motif from the factual scene.",
        "Warning Screens":"Prefer a close laptop or desktop with multiple translucent red warning triangles and a clear compromised-system state; people are optional.",
        "Mobile Alerts":"Prefer a close smartphone in hand showing a bold warning-state interface, with a softly blurred workstation behind it.",
        "System Errors and Updates":"Prefer a realistic laptop or device showing an error, failed update, urgent security patch, or remediation state supported by the scene.",
        "Concerned User":"Prefer a realistic concerned office worker reacting to a phone or laptop warning, using red risk light only around the affected device.",
        "Anonymous Hacker":"Prefer an anonymous hooded silhouette at a laptop with dark navy/cyan lighting, only when an attacker, intrusion, phishing or compromise is explicitly supported.",
        "Mixed Cyber Threats":"Select one distinct motif appropriate to this scene: red warning screen, mobile alert, system error/update, concerned user, or anonymous attacker. Do not repeat the motif used by another asset.",
    }
    return f'''Create a vertical 9:16 cinematic B-roll shot for an Arabic cybersecurity alert.
Visual style: {styles.get(selected, styles["Cinematic AI"])}.
Required setting: {scene.get("visualSetting") or "a real-world location directly relevant to the scene"}.
Visual mode: {scene.get("visualMode") or "environment"}.
Scene: {scene.get("visualSuggestion") or scene.get("onScreenText")}.
Preferred threat editorial motif: {threat_styles.get(req.threat_visual_style, threat_styles["Auto by Content"])}
Campaign visual language: premium UAE public-awareness film, human-centered documentary realism, natural expressions and everyday actions, refined dark navy-to-teal cinematic color grade, soft practical lighting, shallow depth of field, polished but believable government communication style. Compose the subject in the middle and lower portions while preserving clean, low-detail negative space around the lower-middle text zone. Do not generate any text inside the visual asset; typography is added later by the video renderer.
Choose the environment from the required setting and factual content, not from generic cybersecurity imagery. Possible environments include private offices, open workplaces, meeting rooms, government service areas, reception spaces, home or family settings, remote-work desks, cafés, travel environments, data centers, technical rooms, and SOC facilities. Do not use a SOC, server room, large monitoring wall, or multiple cyber screens unless this exact scene requires it. Do not repeat the setting, camera angle, people arrangement, or main device used in another asset from this video.
If visual mode is screen_capture, create a convincing full-frame or close-up screen-capture-style interface that directly depicts the supported event, such as a phishing email layout, suspicious hyperlink focus, fake login form, unusual login warning, malicious attachment state, malware alert, compromised session, or security update. Make the mechanism understandable through layout, icons, warning colors, cursor position, highlights and interface state. Do not invent personal data, credentials, sender names, domains, brands, CVEs, versions, or unsupported incident details. Any interface copy must be abstract, blurred or non-readable; the separate video overlay supplies the real text.
If visual mode is object_only or device_closeup, show the relevant screen, device, update, patch, warning, email or technical object without any person in frame. If visual mode is hacker_scene, show one anonymous hooded attacker as a face-obscured silhouette behind a laptop in a dark navy and cyan cinematic environment, similar to a serious cybersecurity editorial photograph. Use hacker_scene only when the supplied facts explicitly support an attacker, hacking or compromise; never imply identity or attribution.
Show culturally accurate Emirati people only where people improve the scene. Emirati men must have authentic Gulf/Emirati facial features and wear a pristine white kandura, white ghutra and clearly visible black agal. Emirati women must have calm, dignified Emirati facial features and wear an elegant modest black abaya with a black shayla. Family members and children may appear only for relevant personal or family cybersecurity awareness, in a natural, respectful UAE home context. Keep wardrobe culturally accurate and professional. Include realistic phones, tablets, laptops, office computers, home devices, security screens, or servers only when relevant to this exact scene. Natural human movement, realistic hands, coherent screen glow, subtle camera motion, premium commercial production quality. No dialogue and no generated audio is needed. No readable text, logos, captions, watermarks, distorted faces, extra fingers, panic, weapons, sensational criminal stereotypes, or fantasy interfaces.'''

def _cinematic_still_prompt(scene: dict[str, Any], req: VisualAlertRequest, index: int) -> str:
    directions = ["wide environmental establishing composition", "medium private-office composition", "over-the-shoulder device interaction", "human-centered conversational composition", "close-up action with environmental context", "calm wide closing composition"]
    mode = scene.get("visualMode")
    composition = "anonymous hooded attacker silhouette centered behind a laptop, face fully obscured, dark navy/cyan editorial lighting" if mode == "hacker_scene" else "full-frame realistic screen capture or tight device-screen close-up with a clearly visible incident state and no people" if mode in {"screen_capture", "object_only", "device_closeup"} else directions[index]
    return _veo_prompt(scene, req).replace("Create a vertical 9:16 cinematic B-roll shot", "Create a single vertical 9:16 cinematic editorial still image").replace("Natural human movement", "Natural body posture") + f"\nDistinct still-image composition: {composition}. Match this specific scene only and do not reuse the composition of another scene. Sharp photographic detail, strong foreground-midground-background separation, suitable for subtle cinematic pan and zoom."

def _generate_veo_clip(prompt: str, output_path: Path):
    api_key = os.environ["GEMINI_API_KEY"]
    base = "https://generativelanguage.googleapis.com/v1beta"
    model = os.getenv("GEMINI_VIDEO_MODEL", "veo-3.1-fast-generate-preview")
    headers = {"x-goog-api-key":api_key, "Content-Type":"application/json"}
    def ensure_ok(response: httpx.Response, stage: str):
        if response.is_success: return
        try:
            body = response.json(); detail = (body.get("error") or {}).get("message") or str(body.get("error") or body)
        except Exception:
            detail = response.text[:900]
        raise ValueError(f"فشل Veo أثناء {stage} (HTTP {response.status_code}): {detail[:900]}")
    with httpx.Client(timeout=httpx.Timeout(180.0, connect=20.0), follow_redirects=True) as client:
        created = client.post(f"{base}/models/{model}:predictLongRunning", headers=headers, json={"instances":[{"prompt":prompt}], "parameters":{"aspectRatio":"9:16", "resolution":"720p"}})
        ensure_ok(created, "بدء توليد الفيديو"); operation = created.json()
        operation_name = operation.get("name")
        if not operation_name: raise ValueError("لم يُرجع Gemini معرّف عملية فيديو")
        deadline = time.time() + 720
        while time.time() < deadline:
            status = client.get(f"{base}/{operation_name}", headers=headers)
            ensure_ok(status, "متابعة حالة الفيديو"); data = status.json()
            if data.get("done"):
                if data.get("error"): raise ValueError("فشل توليد لقطة Veo: " + str((data["error"] or {}).get("message") or data["error"])[:900])
                samples = (((data.get("response") or {}).get("generateVideoResponse") or {}).get("generatedSamples") or [])
                url = (((samples[0] if samples else {}).get("video") or {}).get("uri"))
                if not url: raise ValueError("اكتمل Veo دون رابط فيديو صالح")
                video = client.get(url, headers={"x-goog-api-key":api_key})
                ensure_ok(video, "تنزيل الفيديو"); output_path.write_bytes(video.content); return
            time.sleep(10)
    raise ValueError("انتهت مهلة انتظار توليد فيديو Veo")

def _split_veo_clip(source: Path, work_dir: Path) -> list[str]:
    clips = []
    for i, (start, length) in enumerate(((0.0, 2.7), (2.7, 2.7), (5.4, 2.6)), 1):
        target = work_dir / f"veo-segment-{i}.mp4"
        result = subprocess.run(["ffmpeg", "-y", "-ss", str(start), "-i", str(source), "-t", str(length), "-an", "-c:v", "libx264", "-preset", "fast", "-crf", "20", "-pix_fmt", "yuv420p", str(target)], capture_output=True, text=True, timeout=180)
        if result.returncode != 0 or not target.exists(): raise ValueError("تعذر تقسيم مقطع Veo إلى ثلاثة مشاهد")
        clips.append(str(target))
    return clips

def _extract_gemini_audio(payload: Any) -> tuple[bytes, str, int, int] | None:
    """Find audio in both the current Interactions REST schema and legacy responses."""
    def walk(value: Any, inherited: dict[str, Any] | None = None):
        if isinstance(value, dict):
            meta = dict(inherited or {})
            for key in ("mime_type", "mimeType", "sample_rate", "sampleRate", "channels"):
                if value.get(key) is not None:
                    meta[key] = value[key]

            mime = str(meta.get("mime_type") or meta.get("mimeType") or "").lower()
            is_audio = value.get("type") in {"audio", "output_audio"} or mime.startswith("audio/")
            data = value.get("data")
            if is_audio and isinstance(data, str) and data:
                try:
                    decoded = base64.b64decode(data, validate=True)
                except (ValueError, TypeError):
                    decoded = b""
                if decoded:
                    rate = int(meta.get("sample_rate") or meta.get("sampleRate") or 24000)
                    channels = int(meta.get("channels") or 1)
                    return decoded, mime or "audio/l16", rate, channels

            for key in ("output_audio", "inlineData", "inline_data", "audio"):
                if key in value:
                    found = walk(value[key], meta)
                    if found:
                        return found
            for key, child in value.items():
                if key not in {"output_audio", "inlineData", "inline_data", "audio", "data"}:
                    found = walk(child, meta)
                    if found:
                        return found
        elif isinstance(value, list):
            for child in value:
                found = walk(child, inherited)
                if found:
                    return found
        return None

    return walk(payload)

def _gemini_response_shape(payload: Any) -> str:
    if not isinstance(payload, dict):
        return type(payload).__name__
    keys = sorted(str(key) for key in payload.keys())[:12]
    step_types = []
    for step in payload.get("steps", []) if isinstance(payload.get("steps"), list) else []:
        if isinstance(step, dict):
            step_types.append(str(step.get("type") or step.get("role") or "unknown"))
    suffix = f"; step_types={step_types[:12]}" if step_types else ""
    return f"keys={keys}{suffix}"

def _gemini_tts(script: dict[str, Any]) -> tuple[bytes, str, int, int]:
    transcript = "\n".join(s["voiceText"] for s in script["scenes"] if s.get("voiceText"))
    if not transcript: raise ValueError("سيناريو التعليق الصوتي فارغ")
    instruction = f'''Read the following Arabic cybersecurity alert in a natural professional UAE/Emirati speaking style.
Male voice. Calm, confident cybersecurity news presenter. Moderate pace, clear Arabic pronunciation, no exaggerated emotion. Keep English product names and cybersecurity terms clearly pronounced. Read only the transcript, without adding any words.

TRANSCRIPT:
{transcript}'''
    model = os.getenv("GEMINI_TTS_MODEL", "gemini-3.1-flash-tts-preview")
    with httpx.Client(timeout=httpx.Timeout(180.0, connect=20.0)) as client:
        response = client.post(
            "https://generativelanguage.googleapis.com/v1beta/interactions",
            headers={"x-goog-api-key":os.environ["GEMINI_API_KEY"], "Content-Type":"application/json", "Api-Revision":"2026-05-20"},
            json={"model":model, "input":instruction, "response_format":{"type":"audio"}, "generation_config":{"speech_config":[{"voice":os.getenv("GEMINI_TTS_VOICE", "Charon")}]}},
        )
        response.raise_for_status()
        payload = response.json()
    audio = _extract_gemini_audio(payload)
    if not audio:
        raise ValueError("لم يُرجع Gemini بيانات صوتية قابلة للقراءة (" + _gemini_response_shape(payload) + ")")
    return audio

def _write_wav(path: Path, audio: bytes, mime_type: str = "audio/l16", sample_rate: int = 24000, channels: int = 1) -> float:
    if audio.startswith(b"RIFF") or "wav" in mime_type.lower():
        path.write_bytes(audio)
    else:
        with wave.open(str(path), "wb") as wf:
            wf.setnchannels(max(1, channels)); wf.setsampwidth(2); wf.setframerate(max(8000, sample_rate)); wf.writeframes(audio)
    with wave.open(str(path), "rb") as wf:
        return wf.getnframes() / float(wf.getframerate())

def _limit_voice_duration(path: Path, duration: float, maximum: float = 46.0) -> float:
    if duration <= maximum: return duration
    tempo = min(2.0, duration / maximum)
    adjusted = path.with_name("voiceover-adjusted.wav")
    result = subprocess.run(["ffmpeg", "-y", "-i", str(path), "-filter:a", f"atempo={tempo:.5f}", str(adjusted)], capture_output=True, text=True, timeout=120)
    if result.returncode != 0 or not adjusted.exists(): raise ValueError("تعذر ضبط مدة المحتوى قبل خاتمة العلامة التجارية")
    shutil.move(str(adjusted), str(path))
    with wave.open(str(path), "rb") as wf: return wf.getnframes() / float(wf.getframerate())

def _write_corporate_music(path: Path, duration: float):
    """Create a stereo cinematic corporate ambient bed with pad, pulse and digital shimmer."""
    rate = 24000
    seconds = min(60.0, max(18.0, duration + 1.5))
    total = int(rate * seconds)
    progression = [(261.63,329.63,392.00),(220.00,261.63,329.63),(174.61,220.00,261.63),(196.00,246.94,293.66)]
    frames = bytearray()
    for n in range(total):
        t = n / rate; chord = progression[int(t // 4) % len(progression)]
        local = t % .5; pluck = math.exp(-local * 7.5)
        pad = sum(math.sin(2 * math.pi * f * t) for f in chord) / 3
        pulse_note = chord[int((t * 2) % 3)] * 2
        pulse = math.sin(2 * math.pi * pulse_note * t) * pluck
        shimmer = math.sin(2 * math.pi * chord[2] * 2 * t) * (.15 + .1 * math.sin(2 * math.pi * .12 * t))
        intro = min(1.0, t / 1.2); outro = min(1.0, max(0.0, seconds - t) / 1.4)
        bed = (pad * .38 + pulse * .20) * intro * outro
        left = bed + shimmer * .12 * intro * outro
        right_shimmer = math.sin(2 * math.pi * chord[2] * 2 * t + .42) * (.15 + .1 * math.sin(2 * math.pi * .12 * t))
        right = bed * .97 + right_shimmer * .12 * intro * outro
        for sample in (left, right):
            value = max(-32767, min(32767, int(sample * 12500)))
            frames.extend(value.to_bytes(2, "little", signed=True))
    with wave.open(str(path), "wb") as wf:
        wf.setnchannels(2); wf.setsampwidth(2); wf.setframerate(rate); wf.writeframes(frames)

def _normalize_final_video_audio(path: Path):
    """Match the supplied reference videos: about -14 LUFS integrated and -1 dBTP."""
    normalized = path.with_name("visual-alert-normalized.mp4")
    result = subprocess.run(["ffmpeg", "-y", "-i", str(path), "-c:v", "copy", "-af", "loudnorm=I=-14:TP=-1:LRA=6", "-c:a", "aac", "-b:a", "192k", str(normalized)], capture_output=True, text=True, timeout=240)
    if result.returncode != 0 or not normalized.exists(): raise ValueError("تعذر ضبط مستوى الصوت النهائي إلى -14 LUFS")
    shutil.move(str(normalized), str(path))

def _append_brand_outro(path: Path):
    parts = sorted((BASE_DIR / "assets").glob("cyberpulse-outro.b64.part*"))
    if not parts: raise ValueError("ملف خاتمة نبض سيبراني غير موجود")
    outro = path.with_name("cyberpulse-brand-outro.mp4")
    outro.write_bytes(base64.b64decode("".join(part.read_text(encoding="ascii") for part in parts)))
    combined = path.with_name("visual-alert-with-outro.mp4")
    filters = "[0:v]fps=30,scale=1080:1920,setsar=1[v0];[1:v]fps=30,scale=1080:1920,setsar=1[v1];[0:a]aresample=48000,aformat=sample_fmts=fltp:channel_layouts=stereo[a0];[1:a]aresample=48000,aformat=sample_fmts=fltp:channel_layouts=stereo[a1];[v0][a0][v1][a1]concat=n=2:v=1:a=1[v][a]"
    result = subprocess.run(["ffmpeg", "-y", "-i", str(path), "-i", str(outro), "-filter_complex", filters, "-map", "[v]", "-map", "[a]", "-c:v", "libx264", "-preset", "medium", "-crf", "20", "-c:a", "aac", "-b:a", "192k", "-movflags", "+faststart", str(combined)], capture_output=True, text=True, timeout=360)
    if result.returncode != 0 or not combined.exists(): raise ValueError("تعذر دمج خاتمة نبض سيبراني: " + (result.stderr or result.stdout)[-700:])
    shutil.move(str(combined), str(path))

def _fit_scene_durations(script: dict[str, Any], audio_duration: float):
    if audio_duration > 46.5: raise ValueError("تعذر ضبط التعليق الصوتي قبل خاتمة العلامة التجارية")
    scenes = script["scenes"]
    weights = [max(1, len(s.get("voiceText", "").split())) for s in scenes]
    total = sum(weights)
    target = min(47.0, max(18.0, audio_duration + 1.0))
    remaining = target
    for i, scene in enumerate(scenes):
        duration = remaining if i == len(scenes)-1 else max(2.5, target * weights[i] / total)
        scene["duration"] = round(duration, 2); remaining -= duration
    script["estimatedDuration"] = round(target, 2)

def _cleanup_visual_jobs():
    cutoff = time.time() - 3600
    expired = []
    with VISUAL_ALERT_JOBS_LOCK:
        for job_id, job in VISUAL_ALERT_JOBS.items():
            if job.get("created_ts", time.time()) < cutoff: expired.append((job_id, job.get("work_dir")))
        for job_id, _ in expired: VISUAL_ALERT_JOBS.pop(job_id, None)
    for _, folder in expired:
        if folder: shutil.rmtree(folder, ignore_errors=True)

def _run_visual_alert_job(job_id: str, req: VisualAlertRequest):
    work_dir = Path(tempfile.mkdtemp(prefix=f"cyberpulse-alert-{job_id[:8]}-"))
    _visual_job_update(job_id, work_dir=str(work_dir), status="analyzing", progress=12, message="جاري تحليل التنبيه...")
    try:
        script = _visual_script(req)
        _visual_job_update(job_id, status="generating_voice", progress=38, message="جاري إنشاء التعليق الصوتي...", script=script)
        audio, mime_type, sample_rate, channels = _gemini_tts(script)
        audio_path = work_dir / "voiceover.wav"
        duration = _write_wav(audio_path, audio, mime_type, sample_rate, channels)
        duration = _limit_voice_duration(audio_path, duration)
        music_path = work_dir / "inspirational-corporate.wav"
        _write_corporate_music(music_path, duration)
        _fit_scene_durations(script, duration)
        clip_paths, image_paths = [], []
        selected = [script["scenes"][i % len(script["scenes"])] for i in range(6)]
        for video_index in range(req.video_count):
            _visual_job_update(job_id, status="generating_visuals", progress=46+video_index*5, message=f"جاري إنشاء الفيديو {video_index+1} من {req.video_count}...", script=script)
            clip_path = work_dir / f"veo-source-{video_index+1}.mp4"
            try:
                scene = selected[(video_index * 2) % len(selected)]
                _generate_veo_clip(_veo_prompt(scene, req) + f"\nUnique video shot {video_index+1} of {req.video_count}. Use a distinct camera angle and scene composition. Do not repeat any earlier shot.", clip_path)
                clip_paths.append(str(clip_path))
            except Exception as veo_error:
                if "429" not in str(veo_error) and "quota" not in str(veo_error).lower(): raise
                _visual_job_update(job_id, message=f"توقّف Veo عند {len(clip_paths)} فيديو بسبب الحصة؛ سيتم ضبط عدد الصور تلقائيًا.", veo_fallback=str(veo_error)[:500])
                break
        image_indices = (1, 3, 5) if len(clip_paths) == 3 else (1, 2, 3, 4, 5) if len(clip_paths) == 1 else (0, 1, 2, 3, 4, 5)
        for position, i in enumerate(image_indices, 1):
            _visual_job_update(job_id, status="generating_visuals", progress=min(74, 52+position*4), message=f"جاري إنشاء الصورة {position} من {len(image_indices)} حسب محتوى المشهد...", script=script)
            image_b64, _ = generate_nano_banana_image(_cinematic_still_prompt(selected[i], req, i), "9:16")
            image_path = work_dir / f"ai-scene-{position}.jpg"
            image_path.write_bytes(base64.b64decode(image_b64)); image_paths.append(str(image_path))
        asset_order = [f"video:{i+1}" for i in range(len(clip_paths))] + [f"image:{i+1}" for i in range(len(image_paths))]
        _visual_job_update(job_id, status="ready_for_review", progress=78, message="المواد جاهزة للمراجعة والترتيب. وافق عليها لبدء الدمج.", script=script, clip_count=len(clip_paths), image_count=len(image_paths), asset_order=asset_order, preview_ready=True)
    except Exception as exc:
        _visual_job_update(job_id, status="failed", message=str(exc)[:1400], detail=str(exc)[:1400], completed_at=datetime.now(timezone.utc).isoformat())

def _regenerate_visual_alert_audio(job_id: str, req: VisualAlertRequest):
    with VISUAL_ALERT_JOBS_LOCK:
        job = VISUAL_ALERT_JOBS.get(job_id)
        if not job: return
        work_dir = Path(job["work_dir"])
    try:
        script = _visual_script(req)
        _visual_job_update(job_id, status="regenerating_audio", progress=42, message="جاري إعادة إنشاء التعليق الصوتي من المحتوى المعدّل...", script=script)
        audio, mime_type, sample_rate, channels = _gemini_tts(script)
        new_audio = work_dir / "voiceover-new.wav"
        duration = _write_wav(new_audio, audio, mime_type, sample_rate, channels)
        duration = _limit_voice_duration(new_audio, duration)
        new_music = work_dir / "inspirational-corporate-new.wav"
        _write_corporate_music(new_music, duration)
        _fit_scene_durations(script, duration)
        shutil.move(str(new_audio), str(work_dir / "voiceover.wav"))
        shutil.move(str(new_music), str(work_dir / "inspirational-corporate.wav"))
        _visual_job_update(job_id, status="ready_for_review", progress=78, message="تم تحديث النص والتعليق الصوتي. الصور والفيديوهات لم تتغير.", script=script, audio_regenerated=True, audio_regeneration_error=None)
    except Exception as exc:
        _visual_job_update(job_id, status="ready_for_review", progress=78, message="تعذر إعادة توليد الصوت؛ بقي الصوت السابق متاحًا.", audio_regeneration_error=str(exc)[:900])

def _render_approved_visual_alert(job_id: str):
    try:
        with VISUAL_ALERT_JOBS_LOCK:
            job = VISUAL_ALERT_JOBS.get(job_id)
            if not job: return
            work_dir = Path(job["work_dir"]); script = job["script"]; asset_order = job.get("asset_order", []); motion_overlays = bool(job.get("motion_overlays")); threat_visual_style = job.get("threat_visual_style", "Auto by Content")
        audio_path = work_dir / "voiceover.wav"
        music_path = work_dir / "inspirational-corporate.wav"
        clip_paths = [str(path) for path in sorted(work_dir.glob("veo-source-*.mp4"))]
        image_paths = [str(path) for path in sorted(work_dir.glob("ai-scene-*.jpg"))]
        asset_lookup = {**{f"video:{i+1}":{"type":"video","path":path} for i,path in enumerate(clip_paths)}, **{f"image:{i+1}":{"type":"image","path":path} for i,path in enumerate(image_paths)}}
        visual_assets = [asset_lookup[token] for token in asset_order if token in asset_lookup]
        if len(visual_assets) != len(asset_lookup): raise ValueError("ترتيب المواد غير مكتمل")
        _visual_job_update(job_id, status="rendering", progress=84, message="تمت الموافقة؛ جاري دمج الفيديو والصور والصوت...")
        props_path = work_dir / "props.json"; output_path = work_dir / "visual-alert.mp4"
        props_path.write_text(json.dumps({"script":script, "audioPath":str(audio_path), "musicPath":str(music_path), "visualAssets":visual_assets, "motionOverlays":motion_overlays, "threatVisualStyle":threat_visual_style}, ensure_ascii=False), encoding="utf-8")
        result = subprocess.run(["node", str(BASE_DIR / "remotion" / "render.mjs"), str(props_path), str(output_path)], cwd=BASE_DIR, capture_output=True, text=True, timeout=600)
        if result.returncode != 0: raise ValueError("فشل Remotion: " + (result.stderr or result.stdout)[-1200:])
        _visual_job_update(job_id, status="rendering", progress=94, message="جاري إضافة خاتمة نبض سيبراني...")
        _append_brand_outro(output_path)
        _visual_job_update(job_id, status="rendering", progress=97, message="جاري ضبط الموسيقى والصوت إلى المستوى المرجعي...")
        _normalize_final_video_audio(output_path)
        _visual_job_update(job_id, status="completed", progress=100, message="اكتمل الفيديو", video_ready=True, completed_at=datetime.now(timezone.utc).isoformat())
    except Exception as exc:
        _visual_job_update(job_id, status="failed", message=str(exc)[:1400], detail=str(exc)[:1400], completed_at=datetime.now(timezone.utc).isoformat())

@app.post("/api/visual-alert/render", status_code=202)
def create_visual_alert(req: VisualAlertRequest):
    if not os.getenv("OPENAI_API_KEY"): raise HTTPException(400, "OPENAI_API_KEY is not configured")
    if not os.getenv("GEMINI_API_KEY"): raise HTTPException(400, "GEMINI_API_KEY is not configured")
    _cleanup_visual_jobs()
    job_id = str(uuid.uuid4())
    with VISUAL_ALERT_JOBS_LOCK:
        VISUAL_ALERT_JOBS[job_id] = {"id":job_id, "status":"pending", "progress":3, "message":"تم إنشاء المهمة", "threat_visual_style":req.threat_visual_style, "created_ts":time.time(), "created_at":datetime.now(timezone.utc).isoformat()}
    threading.Thread(target=_run_visual_alert_job, args=(job_id, req), daemon=True).start()
    return {"id":job_id, "status":"pending", "progress":3, "message":"تم إنشاء المهمة"}

@app.get("/api/visual-alert/status/{job_id}")
def visual_alert_status(job_id: str):
    with VISUAL_ALERT_JOBS_LOCK:
        job = VISUAL_ALERT_JOBS.get(job_id)
        if not job: raise HTTPException(404, "المهمة غير موجودة أو انتهت صلاحيتها")
        return {k:v for k,v in job.items() if k not in {"work_dir", "created_ts"}}

@app.post("/api/visual-alert/regenerate-audio/{job_id}", status_code=202)
def regenerate_visual_alert_audio(job_id: str, req: VisualAlertRequest):
    with VISUAL_ALERT_JOBS_LOCK:
        job = VISUAL_ALERT_JOBS.get(job_id)
        if not job: raise HTTPException(404, "المهمة غير موجودة أو انتهت صلاحيتها")
        if job.get("status") != "ready_for_review": raise HTTPException(409, "انتظر حتى تصبح المواد جاهزة للمراجعة")
        job.update(status="regenerating_audio", progress=35, message="تم استلام المحتوى المعدّل لإعادة توليد الصوت")
    threading.Thread(target=_regenerate_visual_alert_audio, args=(job_id, req), daemon=True).start()
    return {"id":job_id, "status":"regenerating_audio", "progress":35}

@app.post("/api/visual-alert/approve/{job_id}", status_code=202)
def approve_visual_alert(job_id: str, approval: VisualApprovalRequest):
    with VISUAL_ALERT_JOBS_LOCK:
        job = VISUAL_ALERT_JOBS.get(job_id)
        if not job: raise HTTPException(404, "المهمة غير موجودة أو انتهت صلاحيتها")
        if job.get("status") != "ready_for_review": raise HTTPException(409, "المواد ليست جاهزة للموافقة")
        expected = set(job.get("asset_order", []))
        supplied = approval.asset_order
        if len(supplied) != len(expected) or set(supplied) != expected: raise HTTPException(400, "ترتيب المواد غير صالح أو غير مكتمل")
        job.update(status="approval_received", progress=80, message="تم استلام الموافقة", asset_order=supplied, motion_overlays=approval.motion_overlays)
    threading.Thread(target=_render_approved_visual_alert, args=(job_id,), daemon=True).start()
    return {"id":job_id, "status":"approval_received", "progress":80}

@app.get("/api/visual-alert/preview-video/{job_id}/{video_number}")
def visual_alert_preview_video(job_id: str, video_number: int):
    if video_number not in {1, 2, 3}: raise HTTPException(404, "رقم الفيديو غير صالح")
    with VISUAL_ALERT_JOBS_LOCK: job = VISUAL_ALERT_JOBS.get(job_id)
    if not job or job.get("status") not in {"ready_for_review", "approval_received", "rendering", "completed"}: raise HTTPException(404, "معاينة الفيديو غير جاهزة")
    path = Path(job["work_dir"]) / f"veo-source-{video_number}.mp4"
    if not path.exists(): raise HTTPException(404, "لا توجد لقطة فيديو بسبب حدود Veo الحالية")
    return FileResponse(path, media_type="video/mp4", filename=f"preview-{job_id[:8]}.mp4")

@app.get("/api/visual-alert/preview-audio/{job_id}")
def visual_alert_preview_audio(job_id: str):
    with VISUAL_ALERT_JOBS_LOCK: job = VISUAL_ALERT_JOBS.get(job_id)
    if not job or job.get("status") not in {"ready_for_review", "approval_received", "rendering", "completed"}: raise HTTPException(404, "معاينة الصوت غير جاهزة")
    path = Path(job["work_dir"]) / "voiceover.wav"
    if not path.exists(): raise HTTPException(404, "ملف الصوت غير موجود")
    return FileResponse(path, media_type="audio/wav", filename=f"voiceover-{job_id[:8]}.wav")

@app.get("/api/visual-alert/preview-image/{job_id}/{image_number}")
def visual_alert_preview_image(job_id: str, image_number: int):
    if image_number < 1 or image_number > 6: raise HTTPException(404, "رقم الصورة غير صالح")
    with VISUAL_ALERT_JOBS_LOCK: job = VISUAL_ALERT_JOBS.get(job_id)
    if not job or job.get("status") not in {"ready_for_review", "approval_received", "rendering", "completed"}: raise HTTPException(404, "معاينة الصورة غير جاهزة")
    path = Path(job["work_dir"]) / f"ai-scene-{image_number}.jpg"
    if not path.exists(): raise HTTPException(404, "الصورة غير موجودة")
    return FileResponse(path, media_type="image/jpeg")

@app.get("/api/visual-alert/video/{job_id}")
def visual_alert_video(job_id: str):
    with VISUAL_ALERT_JOBS_LOCK: job = VISUAL_ALERT_JOBS.get(job_id)
    if not job or job.get("status") != "completed": raise HTTPException(404, "الفيديو غير جاهز")
    path = Path(job["work_dir"]) / "visual-alert.mp4"
    if not path.exists(): raise HTTPException(404, "انتهت صلاحية ملف الفيديو")
    return FileResponse(path, media_type="video/mp4", filename=f"cyberpulse-alert-{job_id[:8]}.mp4")

def _google_drive_access_token() -> str:
    required = {key:os.getenv(key, "").strip() for key in ("GOOGLE_DRIVE_CLIENT_ID", "GOOGLE_DRIVE_CLIENT_SECRET", "GOOGLE_DRIVE_REFRESH_TOKEN")}
    missing = [key for key, value in required.items() if not value]
    if missing: raise ValueError("إعداد Google Drive غير مكتمل: " + ", ".join(missing))
    with httpx.Client(timeout=30.0) as client:
        response = client.post("https://oauth2.googleapis.com/token", data={"client_id":required["GOOGLE_DRIVE_CLIENT_ID"], "client_secret":required["GOOGLE_DRIVE_CLIENT_SECRET"], "refresh_token":required["GOOGLE_DRIVE_REFRESH_TOKEN"], "grant_type":"refresh_token"})
        if not response.is_success: raise ValueError(f"تعذر تفويض Google Drive (HTTP {response.status_code}): {response.text[:500]}")
        token = str(response.json().get("access_token", ""))
        if not token: raise ValueError("لم يُرجع Google رمز وصول صالحًا")
        return token

def _google_drive_upload(name: str, mime_type: str, payload: bytes, access_token: str) -> dict[str, str]:
    metadata: dict[str, Any] = {"name":name}
    folder_id = os.getenv("GOOGLE_DRIVE_FOLDER_ID", "").strip()
    if folder_id: metadata["parents"] = [folder_id]
    headers = {"Authorization":f"Bearer {access_token}", "Content-Type":"application/json; charset=UTF-8", "X-Upload-Content-Type":mime_type, "X-Upload-Content-Length":str(len(payload))}
    with httpx.Client(timeout=httpx.Timeout(300.0, connect=30.0), follow_redirects=False) as client:
        start = client.post("https://www.googleapis.com/upload/drive/v3/files?uploadType=resumable&supportsAllDrives=true&fields=id,name,webViewLink", headers=headers, json=metadata)
        if not start.is_success or not start.headers.get("Location"): raise ValueError(f"تعذر بدء رفع {name} إلى Google Drive (HTTP {start.status_code}): {start.text[:500]}")
        uploaded = client.put(start.headers["Location"], headers={"Content-Type":mime_type, "Content-Length":str(len(payload))}, content=payload)
        if not uploaded.is_success: raise ValueError(f"تعذر رفع {name} إلى Google Drive (HTTP {uploaded.status_code}): {uploaded.text[:500]}")
        data = uploaded.json()
        return {"id":str(data.get("id", "")), "url":str(data.get("webViewLink", "")), "name":str(data.get("name", name))}

@app.get("/api/visual-alert/archives")
def visual_alert_archives():
    with db_conn() as c:
        return c.execute("SELECT id,title,content,required_action,formatted_content,drive_video_url,drive_text_url,created_at FROM visual_alert_archives ORDER BY created_at DESC LIMIT 100").fetchall()

@app.post("/api/visual-alert/save/{job_id}")
def save_visual_alert(job_id: str, req: VisualSaveRequest):
    with VISUAL_ALERT_JOBS_LOCK: job = VISUAL_ALERT_JOBS.get(job_id)
    if not job or job.get("status") != "completed": raise HTTPException(404, "الفيديو غير جاهز أو انتهت صلاحية المهمة")
    video_path = Path(job["work_dir"]) / "visual-alert.mp4"
    if not video_path.exists(): raise HTTPException(404, "انتهت صلاحية ملف الفيديو")
    formatted = req.formatted_content.strip() or f"{req.title}\n\n{req.content}\n\nالإجراء المطلوب:\n{req.required_action}"
    archive_id = str(job.get("archive_id") or uuid.uuid4())
    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "-", req.title).strip("-")[:60] or f"cyberpulse-{job_id[:8]}"
    with db_conn() as c:
        c.execute("INSERT INTO visual_alert_archives(id,title,content,required_action,formatted_content) VALUES(%s,%s,%s,%s,%s) ON CONFLICT(id) DO UPDATE SET title=EXCLUDED.title,content=EXCLUDED.content,required_action=EXCLUDED.required_action,formatted_content=EXCLUDED.formatted_content", (archive_id,req.title,req.content,req.required_action,formatted)); c.commit()
    _visual_job_update(job_id, archive_id=archive_id, saved_to_site=True)
    try:
        token = _google_drive_access_token()
        video = _google_drive_upload(f"{safe_name}.mp4", "video/mp4", video_path.read_bytes(), token)
        text_file = _google_drive_upload(f"{safe_name}-content.txt", "text/plain; charset=UTF-8", formatted.encode("utf-8"), token)
        with db_conn() as c:
            c.execute("UPDATE visual_alert_archives SET drive_video_id=%s,drive_video_url=%s,drive_text_id=%s,drive_text_url=%s WHERE id=%s", (video["id"],video["url"],text_file["id"],text_file["url"],archive_id)); c.commit()
        _visual_job_update(job_id, saved_to_drive=True, archive_id=archive_id, drive_video_url=video["url"], drive_text_url=text_file["url"])
        return {"saved":True, "saved_to_site":True, "saved_to_drive":True, "archive_id":archive_id, "drive_video_url":video["url"], "drive_text_url":text_file["url"]}
    except Exception as exc:
        return {"saved":True, "saved_to_site":True, "saved_to_drive":False, "archive_id":archive_id, "warning":f"تم حفظ النص في الموقع، لكن تعذر الحفظ في Google Drive: {str(exc)[:800]}"}

def _video_url(output: Any) -> str:
    if isinstance(output, str):
        return output
    if isinstance(output, list) and output:
        return _video_url(output[0])
    if isinstance(output, dict):
        for key in ("url", "video_url", "video", "output"):
            if output.get(key):
                return _video_url(output[key])
    raise ValueError("Bytez returned no playable video URL")

def _available_bytez_video_models(client: httpx.Client) -> list[str]:
    url = "https://api.bytez.com/models/v2/list/models"
    headers = {"Authorization": os.environ["BYTEZ_API_KEY"]}
    response = client.get(url, headers=headers, params={"task": "text-to-video"})
    if response.status_code >= 500:
        response = client.get(url, headers=headers, params={"task": "text-to-video"})
    if response.status_code >= 500:
        response = client.get(url, headers=headers)
    response.raise_for_status()
    payload = response.json()
    if isinstance(payload, dict) and payload.get("error"):
        raise ValueError(str(payload["error"]))
    rows = payload.get("output", payload.get("models", [])) if isinstance(payload, dict) else payload
    if isinstance(rows, dict):
        rows = rows.get("models", rows.get("items", rows.get("data", [])))
    if not isinstance(rows, list):
        rows = []
    models = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        task = row.get("task", row.get("tasks", ""))
        tasks = task if isinstance(task, list) else [task]
        model_id = row.get("modelId", row.get("id", row.get("model")))
        if model_id and "text-to-video" in tasks:
            models.append(str(model_id).strip())
    return models

UNAVAILABLE_BYTEZ_VIDEO_MODELS = {"ali-vilab/text-to-video-ms-1.7b"}

def _bytez_video_candidates(client: httpx.Client) -> list[str]:
    configured = os.getenv("BYTEZ_VIDEO_MODEL", "").strip()
    models = _available_bytez_video_models(client)
    configured_is_usable = configured and configured not in UNAVAILABLE_BYTEZ_VIDEO_MODELS
    candidates = ([configured] if configured_is_usable else []) + models
    unique = []
    for model in candidates:
        if model and model not in UNAVAILABLE_BYTEZ_VIDEO_MODELS and model not in unique:
            unique.append(model)
    if unique:
        return unique
    raise ValueError(
        "لم يعثر Bytez على أي نموذج text-to-video متاح لهذا المفتاح. "
        "تحقق من توفر نماذج الفيديو والرصيد في حساب Bytez."
    )

def _run_video_job(job_id: str, req: NewsVideoRequest):
    model = os.getenv("BYTEZ_VIDEO_MODEL", "automatic").strip() or "automatic"
    prompt = f'''Create a premium vertical 9:16 editorial cybersecurity video for CYBER PULSE.
Topic: {req.headline}
Factual context: {req.summary}
Threat category: {req.threat_type}
Visual direction: {req.visual_brief or "abstract digital security environment"}
Style: {req.style}. Target duration: approximately {req.duration} seconds.
SEMANTIC RULE: show the affected technology and reported action/risk directly. Every prominent object must be traceable to the topic. Never replace a technical term with an unrelated physical metaphor. Software patches and updates must look like software update/patch deployment, not plumbing, water leaks, bandages, construction, or hand tools. Simplified non-readable interfaces are allowed only when they clarify the story.
Use dark navy and black with cyan and cyber blue highlights. Use elegant cinematic movement, realistic lighting, and one coherent scene. Do not show readable text, letters, numbers, captions, logos, watermarks, unrelated dashboards, or distorted interfaces. Do not depict graphic violence or panic.'''
    try:
        with httpx.Client(timeout=httpx.Timeout(300.0, connect=20.0)) as client:
            payload = None
            failures = []
            for model in _bytez_video_candidates(client):
                response = client.post(
                    f"https://api.bytez.com/models/v2/{model}",
                    headers={"Authorization": os.environ["BYTEZ_API_KEY"], "Content-Type":"application/json"},
                    json={"text": prompt},
                )
                if response.status_code == 404:
                    failures.append(model)
                    continue
                response.raise_for_status()
                payload = response.json()
                break
            if payload is None:
                raise ValueError(
                    "نماذج Bytez التالية غير متاحة (404): " + ", ".join(failures)
                    + ". اختر نموذج text-to-video متاحًا من لوحة Bytez."
                )
        error = payload.get("error") if isinstance(payload, dict) else None
        if error: raise ValueError(str(error))
        output = payload.get("output", payload) if isinstance(payload, dict) else payload
        result = {"status":"completed", "video_url":_video_url(output), "model":model, "completed_at":datetime.now(timezone.utc).isoformat()}
    except Exception as exc:
        result = {"status":"failed", "detail":str(exc)[:1000], "model":model, "completed_at":datetime.now(timezone.utc).isoformat()}
    with VIDEO_JOBS_LOCK:
        VIDEO_JOBS[job_id].update(result)

@app.post("/api/news-video", status_code=202)
def create_news_video(req: NewsVideoRequest):
    if not os.getenv("BYTEZ_API_KEY"): raise HTTPException(400, "BYTEZ_API_KEY is not configured")
    job_id = str(uuid.uuid4())
    with VIDEO_JOBS_LOCK:
        VIDEO_JOBS[job_id] = {"id":job_id, "status":"processing", "created_at":datetime.now(timezone.utc).isoformat()}
    threading.Thread(target=_run_video_job, args=(job_id, req), daemon=True).start()
    return VIDEO_JOBS[job_id]

@app.get("/api/news-video/{job_id}")
def news_video_status(job_id: str):
    with VIDEO_JOBS_LOCK:
        job = VIDEO_JOBS.get(job_id)
        if not job: raise HTTPException(404, "Video job not found or the service was restarted")
        return dict(job)

@app.get("/api/news-sources")
def news_sources(): return {"count":len(load_cyber_sources()), "sources":load_cyber_sources()}

@app.post("/api/search-news")
def search_news(req: NewsSearchRequest):
    if not os.getenv("OPENAI_API_KEY"): raise HTTPException(400, "OPENAI_API_KEY is not configured")
    sources = load_cyber_sources()
    if not sources: raise HTTPException(500, "Cybersecurity source list is not available")
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    source_lines = "\n".join(f"- {s['name']}: {s['url']}" for s in sources)
    prompt = f'''You are the live news discovery engine for the Arabic cybersecurity publication "نبض سيبراني | CYBER PULSE".
Today is {today}. Search the web for the newest meaningful cybersecurity news published within the last {req.days} days.

STRICT SOURCE POLICY:
Use ONLY the approved source list below. Do not return an article from any other domain. Prefer primary official advisories and vendor research when they are the original source; use established cybersecurity media for major breaking stories. Exclude generic evergreen pages, old articles, duplicate syndications, opinion-only pieces, and items without a publication date.

APPROVED SOURCES:
{source_lines}

Select up to {req.limit} distinct, high-value items relevant to cybersecurity professionals and government/enterprise organizations. Prioritize: active exploitation, significant vulnerabilities, breaches/incidents, ransomware, threat campaigns, identity/cloud/security platform issues, critical infrastructure, major regulatory/operational cyber developments.

Return ONLY valid JSON with this exact shape:
{{"items":[{{
"title_ar":"concise Arabic headline",
"title_original":"original article/advisory title",
"date":"exact publication date from source",
"content_type":"خبر or ثغرة أمنية or تنبيه أمني or اختراق/تسريب or بحث أمني",
"severity":"حرج or عالي or متوسط or منخفض or empty if the source does not state it",
"cve":"exact CVE identifier(s) or empty",
"source":"approved source name",
"source_url":"approved source home/advisory URL",
"url":"direct URL to the exact article/advisory on an approved domain",
"summary_ar":"2-3 factual Arabic sentences",
"news_text_ar":"a self-contained factual Arabic news text of 100-220 words based only on the source, preserving date, affected product/sector, attack/vulnerability details and impact; do not invent recommendations",
"recommendations":["only recommendations explicitly stated by the source"],
"relevance":"one short Arabic sentence explaining why this matters"
}}]}}
Do not fabricate a CVE, severity, date, recommendation, or URL.'''
    try:
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        model = os.getenv("OPENAI_MODEL", "gpt-5")
        response = client.responses.create(model=model, input=prompt, tools=[{"type":"web_search"}], store=False)
        data = extract_json(response.output_text)
        valid, seen = [], set()
        for item in data.get("items", []):
            url = str(item.get("url", "")).strip()
            key = (url.lower(), str(item.get("title_original", item.get("title_ar", ""))).strip().lower())
            if not url or not url_is_approved(url) or key in seen: continue
            seen.add(key)
            valid.append(item)
            if len(valid) >= req.limit: break
        return {"searched_at":datetime.now(timezone.utc).isoformat(), "days":req.days, "source_count":len(sources), "items":valid}
    except Exception as e:
        raise HTTPException(500, f"News search failed: {e}")

@app.get("/api/archive")
def archive_list():
    with db_conn() as c: return c.execute("SELECT id,topic_id,topic,domain,post_type,platform,content,created_at,updated_at FROM posts ORDER BY created_at DESC").fetchall()

@app.get("/api/archive/used-topic-ids")
def used_ids():
    with db_conn() as c: return {"topic_ids":[r["topic_id"] for r in c.execute("SELECT DISTINCT topic_id FROM posts WHERE topic_id IS NOT NULL").fetchall()]}

@app.get("/api/archive/{post_id}")
def archive_get(post_id: str):
    with db_conn() as c:
        post = c.execute("SELECT id,topic_id,topic,domain,post_type,platform,content,created_at,updated_at FROM posts WHERE id=%s", (post_id,)).fetchone()
    if not post: raise HTTPException(404, "المنشور غير موجود")
    return post

@app.post("/api/archive")
def archive_save(p: ArchivePost):
    pid = p.id or str(uuid.uuid4()); now = datetime.now(timezone.utc)
    with db_conn() as c:
        c.execute("INSERT INTO posts(id,topic_id,topic,domain,post_type,platform,content,created_at,updated_at) VALUES(%s,%s,%s,%s,%s,%s,%s::jsonb,%s,%s) ON CONFLICT(id) DO UPDATE SET content=EXCLUDED.content,updated_at=EXCLUDED.updated_at", (pid,p.topic_id,p.topic,p.domain,p.post_type,p.platform,json.dumps(p.content,ensure_ascii=False),now,now)); c.commit()
    return {"id":pid,"saved":True}

@app.delete("/api/archive/{post_id}")
def archive_delete(post_id: str):
    with db_conn() as c:
        cur = c.execute("DELETE FROM posts WHERE id=%s", (post_id,)); c.commit(); return {"deleted":cur.rowcount>0}

def demo_payload(req):
    n = req.slides if req.post_type == "Carousel" else 1
    return {"mode":"demo","title":req.topic,"hook":req.topic,"caption":"محتوى تجريبي.","recommendations":[],"cta":"شارك رأيك.","keywords":["GRC"],"hashtags":["#GRC"],"slides":[{"number":i+1,"headline":req.topic,"body":"نص تجريبي."} for i in range(n)],"sources":[]}

@app.post("/api/generate-content")
def generate_content(req: ContentRequest):
    if not os.getenv("OPENAI_API_KEY"): return demo_payload(req)
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY")); model = os.getenv("OPENAI_MODEL", "gpt-5")
    reference = linkedin_reference_for(req.reference_week, req.reference_slot)
    grounding = ""
    source_payload = []
    if reference:
        item, folder_url = reference
        source = item.get("source", {})
        source_payload = [{
            "name": source.get("name", ""),
            "pages": source.get("pages", ""),
            "url": folder_url,
            "why_relevant": item.get("grounding", "")
        }]
        grounding = f"""
Use this supplied reference as the factual foundation. Do not add facts, figures, standards claims, or page citations that are not supported by it.
Reference: {source.get('name', '')}, pages {source.get('pages', '')}.
Grounding notes: {item.get('grounding', '')}
The caption should paraphrase the reference in an original voice and must end with a short source line naming the reference and pages. Do not quote long passages.
"""
    prompt = f"Create publish-ready Arabic {req.post_type} about {req.topic} for government/enterprise cybersecurity professionals. Exactly {req.slides} slides if carousel. Each slide headline MUST be concise: maximum 9 words and maximum 2 visual lines. Each slide body MUST be maximum 32 words, written as one compact idea suitable for no more than 4 visual lines. Put extended explanations in the caption, never in slide body. Return ONLY JSON with title,hook,caption,recommendations,cta,keywords,hashtags,slides(number,headline,body),sources. Never invent citations. Hashtags never belong in slides. {grounding}"
    kw = {"model":model,"input":prompt,"store":False}
    if req.use_web_search: kw["tools"]=[{"type":"web_search"}]
    try:
        d = extract_json(client.responses.create(**kw).output_text)
        for slide in d.get("slides", []):
            for field, limit in (("headline", 9), ("body", 32)):
                words = str(slide.get(field, "")).split()
                slide[field] = " ".join(words[:limit]) + ("…" if len(words) > limit else "")
        if source_payload:
            d["sources"] = source_payload
            d["source_grounded"] = True
        d["mode"]="openai"; return d
    except Exception as e: raise HTTPException(500, f"Generation failed: {e}")

@app.post("/api/extract-news-file")
async def extract_news_file(request: Request):
    """Extract an alert entirely in memory; the uploaded bytes are never persisted or archived."""
    max_bytes = 8 * 1024 * 1024
    raw = await request.body()
    if not raw: raise HTTPException(400, "الملف فارغ")
    if len(raw) > max_bytes: raise HTTPException(413, "حجم الملف يتجاوز 8 ميجابايت")
    filename = unquote(request.headers.get("x-news-filename", "alert.txt"))
    suffix = Path(filename).suffix.lower()
    allowed = {".pdf", ".docx", ".txt", ".md", ".csv", ".json"}
    if suffix not in allowed: raise HTTPException(415, "نوع الملف غير مدعوم. استخدم PDF أو DOCX أو TXT أو MD أو CSV أو JSON")
    try:
        if suffix == ".pdf":
            if not raw.startswith(b"%PDF"): raise ValueError("توقيع PDF غير صالح")
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        elif suffix == ".docx":
            if not zipfile.is_zipfile(io.BytesIO(raw)): raise ValueError("توقيع DOCX غير صالح")
            from docx import Document
            document = Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in document.paragraphs)
            for table in document.tables:
                text += "\n" + "\n".join(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        else:
            text = raw.decode("utf-8-sig")
    except UnicodeDecodeError:
        raise HTTPException(422, "تعذر قراءة ترميز الملف النصي؛ احفظه بترميز UTF-8")
    except Exception as exc:
        raise HTTPException(422, f"تعذر استخراج النص من الملف: {exc}")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) < 10: raise HTTPException(422, "لم يُعثر على نص قابل للقراءة؛ قد يكون الملف صورة ممسوحة ضوئيًا")
    text = text[:12000]
    first_line = next((line.strip() for line in text.splitlines() if len(line.strip()) >= 3), "")
    suggested_title = first_line[:300] if len(first_line) <= 300 else Path(filename).stem[:300]
    return {"filename":Path(filename).name, "title":suggested_title, "text":text, "stored":False}

@app.post("/api/parse-news")
def parse_news(req: NewsParseRequest):
    if not os.getenv("OPENAI_API_KEY"): raise HTTPException(400, "OPENAI_API_KEY is not configured")
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY")); model = os.getenv("OPENAI_MODEL", "gpt-5")
    prompt = f'''You are the editorial intelligence engine for the Arabic cybersecurity publication "نبض سيبراني | CYBER PULSE".
Parse ONLY supplied facts. Never invent or silently correct CVEs, dates, severity, vendors, sources or recommendations.
HEADLINE:\n{req.title}\n\nPASTED NEWS:\n{req.news}
Find the exact original article/advisory page on the named source website when possible. The date field MUST be the publication date displayed by that source page, never today's date, the date of analysis, or an inferred date. Return an empty date if the original publication date cannot be verified.
Return ONLY valid JSON with: headline, severity (حرج/عالي/متوسط/منخفض/or empty), date, date_verified (boolean), source_url (direct exact article URL or empty), cve, summary (2-3 concise Arabic sentences), recommendations (only supplied), source, entities, threat_type, visual_brief, caption, hashtags.
For visual_brief, describe ONE direct, literal editorial scene that immediately identifies the affected technology, vendor/product category, and the reported action or risk. Prefer recognizable product geometry, device/server context, patch/update objects, cloud/email/browser/mobile cues, or incident-specific objects that are explicitly supported by the supplied news. Do NOT replace the subject with a loose metaphor. For example, software patches must look like software update/patch deployment, never plumbing, water leaks, bandages, construction, tools, or physical repair. The brief MUST NOT request readable words, interface labels, captions, diagrams, flowcharts, logos containing text, letters or numbers. It may request simplified non-readable interface shapes only when the story is specifically about software, updates, identity, cloud, email, browsers, or mobile apps.
Caption must be one polished, self-contained Arabic post ready to paste into BOTH LinkedIn and Instagram. Use 140-260 words and rely only on supplied facts.
Format the caption as readable social copy using plain text and intentional line breaks:
1. Open with a strong, factual hook or headline (no clickbait).
2. Add a short paragraph explaining what happened and why it matters.
3. Break key facts into 2-5 short lines beginning with 🔹 when the supplied material supports distinct facts such as the affected product/sector, vulnerability or attack method, impact, severity, CVE, or date. Never create a bullet merely to fill the structure.
4. If recommendations were supplied, introduce them with "ما الذي يجب فعله؟" and list each action on a separate line beginning with ✅. If none were supplied, omit this section completely.
5. Add "الخلاصة:" followed by one concise practical takeaway based only on the supplied facts.
6. End with one natural engagement question or call to action, then place @cyberpulse_ar on its own line.
Keep paragraphs short, professional, direct, RTL-friendly, and easy to scan on mobile. Do not use Markdown headings, asterisks, numbered lists, excessive emojis, or hashtags inside caption.
Hashtags must be a JSON array of 6-10 concise Arabic or English hashtags relevant to the supplied story, each beginning with #. Always include #نبض_سيبراني and #الأمن_السيبراني.'''
    try:
        parsed = extract_json(client.responses.create(model=model, input=prompt, tools=[{"type":"web_search"}], store=False).output_text)
        source_url = str(parsed.get("source_url", "")).strip()
        pasted_urls = re.findall(r'https?://[^\s<>"\']+', req.news)
        candidate_urls = pasted_urls + ([source_url] if source_url else [])
        verified_date, verified_url = "", ""
        for candidate in candidate_urls:
            candidate = candidate.rstrip(".,،؛;:!?)]}")
            if not url_is_approved(candidate):
                continue
            verified_date = source_publication_date(candidate)
            if verified_date:
                verified_url = candidate
                break
        if verified_date:
            parsed["date"] = verified_date
            parsed["date_verified"] = True
            parsed["source_url"] = verified_url
        else:
            parsed["date_verified"] = bool(parsed.get("date_verified") and source_url and url_is_approved(source_url))
            if not parsed["date_verified"]:
                parsed["date"] = ""
        return parsed
    except Exception as e: raise HTTPException(500, f"News parsing failed: {e}")

ARABIC_MONTHS = {
    1:"يناير", 2:"فبراير", 3:"مارس", 4:"أبريل", 5:"مايو", 6:"يونيو",
    7:"يوليو", 8:"أغسطس", 9:"سبتمبر", 10:"أكتوبر", 11:"نوفمبر", 12:"ديسمبر",
}

def _format_source_date(value: str) -> str:
    raw = html.unescape(str(value or "")).strip()
    if not raw:
        return ""
    iso = raw.replace("Z", "+00:00")
    try:
        dt = datetime.fromisoformat(iso)
        return f"{dt.day} {ARABIC_MONTHS[dt.month]} {dt.year}"
    except ValueError:
        pass
    match = re.search(r'\b(20\d{2})-(\d{1,2})-(\d{1,2})\b', raw)
    if match:
        year, month, day = map(int, match.groups())
        if month in ARABIC_MONTHS:
            return f"{day} {ARABIC_MONTHS[month]} {year}"
    return raw[:80]

def source_publication_date(url: str) -> str:
    try:
        with httpx.Client(timeout=httpx.Timeout(20.0, connect=10.0), follow_redirects=True) as client:
            response = client.get(url, headers={"User-Agent":"Mozilla/5.0 CyberPulseBot/1.0"})
            response.raise_for_status()
        page = response.text[:2_000_000]
        patterns = [
            r'<meta[^>]+(?:property|name)=["\']article:published_time["\'][^>]+content=["\']([^"\']+)',
            r'<meta[^>]+content=["\']([^"\']+)["\'][^>]+(?:property|name)=["\']article:published_time["\']',
            r'<meta[^>]+(?:property|name)=["\'](?:datePublished|date|pubdate|publish-date)["\'][^>]+content=["\']([^"\']+)',
            r'["\']datePublished["\']\s*:\s*["\']([^"\']+)',
        ]
        for pattern in patterns:
            match = re.search(pattern, page, re.IGNORECASE)
            if match:
                return _format_source_date(match.group(1))
    except Exception:
        return ""
    return ""

def visual_prompt(req: ImageRequest):
    if req.visual_style == "Cyber Pulse":
        story = f"{req.title} {req.body} {req.visual_direction}".lower()
        zoom_contract = ""
        if "zoom" in story or "زووم" in story:
            zoom_contract = '''\nMANDATORY ZOOM ANNOTATION VISUAL CONTRACT: show a recognizable modern video-meeting scene with several participant tiles; a shared-screen canvas; a clearly visible annotation pen/drawing stroke/cursor acting on that shared canvas; and an unauthorized-control path spreading from the annotated shared screen toward two or more participant laptops/devices. Use red only for the hostile control path and cyan/blue for the legitimate meeting. A generic laptop, generic participant grid, large arrow, isolated user icon, update window, download screen, or single-device warning is NOT sufficient. Do not generate readable Zoom text or logos; communicate the platform through its familiar blue video-meeting visual language and meeting layout.\n'''
        variant_contracts = {
            1: "VARIANT 1 — DIRECT CAUSAL SCENE: use one dominant focal point and a clear visual path from the vulnerable technology/action to the impact.",
            2: "VARIANT 2 — SPLIT EDITORIAL SCENE: use a clearly divided but cinematic source-versus-impact or before-versus-after composition, connected by the exact attack/risk mechanism.",
            3: "VARIANT 3 — ENTERPRISE CONTEXT SCENE: place the exact technology and mechanism inside a realistic enterprise environment with people/devices where relevant, while keeping the technical cause and impact unmistakable.",
        }
        variant_contract = variant_contracts[req.variant_index]
        return f'''Create ONLY the visual background artwork for a premium cybersecurity news post. Before composing, identify the exact subject, affected technology/vendor category, event, and risk mechanism from the supplied title, context, and visual direction.
FORMAT: vertical Instagram 4:5 portrait.
{variant_contract}
This composition must be substantially different from the other variants while preserving the same factual story and Cyber Pulse identity.
BRAND VISUAL SYSTEM: deep black/dark navy #050B12, Cyber Blue #0A84FF, Cyan #00D1C7. Red only when the supplied risk is high/critical. Subtle circuit patterns, restrained digital grid, soft blue/cyan atmospheric glow. Premium enterprise cybersecurity media publication quality, sophisticated and minimal, never gaming-like.
SEMANTIC ACCURACY IS THE HIGHEST PRIORITY: represent the news subject directly and literally. Every prominent object must be traceable to the supplied story. Show the affected technology and the reported action/risk in one coherent scene. Do not invent an unrelated visual metaphor. If the story is about software updates or security patches, show an update/patch deployment environment, recognizable software/platform geometry, prioritized critical update objects, protected enterprise devices or servers, and security status cues. Never translate "patch", "leak", "bug", "cloud", "virus", "worm", "gateway", or similar technical terms into unrelated physical objects unless the supplied story is actually about those physical objects.
STRICT COMPOSITION: Build a cinematic explanatory editorial scene, not a generic stock-style cybersecurity image. The scene must visually tell the incident in one glance: affected platform/context -> vulnerable feature or action -> visible impact on the affected devices. Use depth, clear focal hierarchy, realistic enterprise people/devices where relevant, and precise cyan-versus-red visual causality. Concentrate the explanatory action in the center/lower 55-60%. Preserve a genuinely empty, dark, low-detail text-safe editorial panel across the upper 35-40%; no faces, bright objects, interface panels, attack paths, or important details may enter that panel. Keep additional safe space at the top-right for the Cyber Pulse logo.
ABSOLUTE NO-TEXT RULE: ZERO readable Arabic or English, words, letters, numbers, CVEs, product names, UI labels, captions, headlines, hashtags, watermarks, signatures, brand names, pseudo-text or typographic logo marks.
Simplified non-readable interface panels, update progress shapes, product geometry, patch tiles, device screens, and security-status cards are allowed only when they directly clarify the supplied story. They must contain no text or pseudo-text.
DO NOT CREATE flowcharts, generic dashboards unrelated to the story, browser directory lists, dense tables, hacker hoodies, masks, skulls, Matrix code, random binary streams, plumbing, water leaks, bandages, construction tools, repair tools, or clutter unless explicitly required by the factual story.
DIRECT VISUAL BRIEF: {req.visual_direction}
NEWS TITLE: {req.title}
FACTUAL CONTEXT: {req.body}
{zoom_contract}
Final self-check before rendering: (1) would a viewer identify the technology and event without reading the headline, (2) can the viewer follow the attack or risk mechanism from its origin to its impact, and (3) does the image have a clear editorial hierarchy rather than merely containing the right objects? If any answer is no, revise before rendering.'''
    common = f"Artwork only, 4:5 portrait. Concept: {req.title}. Context: {req.body}. ABSOLUTE: zero readable text, letters, numbers, labels, headings, framework names, interface copy, hashtags, watermarks or pseudo-text. Do not create a labeled infographic, poster, slide, diagram with captions, or text-bearing UI. Reserve the top 28% as a clean low-detail typography-safe zone and the bottom 22% as another clean low-detail typography-safe zone. Place the single main visual subject only in the central 50%, fully visible and not cropped. {req.visual_direction}"
    if req.domain.strip().upper() == "GRC": style = "Premium light government/enterprise GRC editorial illustration; white/cool-gray; navy/royal blue; restrained green/orange/red status accents; one strong central symbolic scene using polished enterprise vector/semi-3D objects; no internal labels, no multi-panel infographic, no repeated mini diagrams, generous whitespace, no hacker clichés."
    elif req.visual_style == "Executive Minimal": style = "Executive minimal artwork, white background, navy/blue, strong central metaphor, whitespace."
    else: style = "Structured professional infographic artwork, light background, central concept and restrained supporting visuals."
    return style + "\n" + common

def review_artwork(client: OpenAI, req: ImageRequest, image_b64: str):
    review_prompt = f'''You are the visual quality-control reviewer for the Arabic cybersecurity publication "نبض سيبراني | CYBER PULSE".
Evaluate the supplied generated artwork against the factual news context. Review the IMAGE itself, not merely the prompt.
Use ONLY the supplied news title, factual context, and required visual direction. Never require an object, screen, feature, update window, download, patch, vendor, or attack step that is not supported by this specific story. Do not carry requirements from another cybersecurity story.

NEWS TITLE: {req.title}
FACTUAL CONTEXT: {req.body}
REQUIRED VISUAL DIRECTION: {req.visual_direction}

Score these criteria:
1. The affected technology/vendor/product category is immediately identifiable without headline text.
2. The reported event or attack mechanism is visually clear and technically relevant.
3. Every prominent object is traceable to the supplied story; there are no loose or misleading metaphors.
4. The scene has premium Cyber Pulse editorial quality: dark navy, cyan/teal, restrained risk red, clean and professional.
5. The upper 35-40% remains usable for Arabic headline and metadata, and the top-right has safe logo space.
6. There is no readable generated text, pseudo-text, watermark, distorted typography, hacker hoodie, or unrelated clutter.
7. The image explains a causal story, not merely a collection of relevant objects: the origin/action, path or mechanism, and impact are visually connected.
8. The composition has editorial hierarchy and cinematic depth comparable to a premium technology-news cover: one dominant focal point, supporting context, and no stock-photo/generic-dashboard feeling.

For Zoom Annotation / screen-sharing takeover stories specifically, a strong image should clearly show a video meeting with several participants, a shared-screen canvas, an annotation pen/drawing stroke/cursor acting on that shared screen, and unauthorized control spreading from the annotated screen toward two or more participant devices. A generic laptop, generic meeting grid, arrow, user icon, software-update window, or download screen is insufficient and must not be requested.

Return ONLY valid JSON:
{{"semantic_match":true,"score":0,"technology_visible":true,"mechanism_visible":true,"composition_ok":true,"issues":["concise issue"],"retry_direction":"specific English correction prompt for the image generator","summary_ar":"سطر عربي مختصر يشرح نتيجة المراجعة"}}
Set semantic_match=true only when score is at least 82 and technology_visible, mechanism_visible, and composition_ok are all true. A visually generic image cannot pass even if it contains the correct platform and objects.'''
    response = client.responses.create(
        model=os.getenv("OPENAI_VISION_MODEL", os.getenv("OPENAI_MODEL", "gpt-5")),
        input=[{"role":"user","content":[
            {"type":"input_text","text":review_prompt},
            {"type":"input_image","image_url":f"data:image/jpeg;base64,{image_b64}"},
        ]}],
        store=False,
    )
    review = extract_json(response.output_text)
    score = max(0, min(100, int(review.get("score", 0))))
    review["score"] = score
    review["semantic_match"] = bool(
        review.get("semantic_match") and score >= 82
        and review.get("technology_visible") and review.get("mechanism_visible")
        and review.get("composition_ok")
    )
    return review

def _find_gemini_image_data(value: Any) -> str:
    if isinstance(value, dict):
        output_image = value.get("output_image")
        if isinstance(output_image, dict) and output_image.get("data"):
            return str(output_image["data"])
        if value.get("type") == "image" and value.get("data"):
            return str(value["data"])
        for child in value.values():
            found = _find_gemini_image_data(child)
            if found:
                return found
    elif isinstance(value, list):
        for child in value:
            found = _find_gemini_image_data(child)
            if found:
                return found
    return ""

def generate_nano_banana_image(prompt: str, aspect_ratio: str = "4:5") -> tuple[str, str]:
    api_key = os.getenv("GEMINI_API_KEY", "").strip()
    if not api_key:
        raise ValueError("GEMINI_API_KEY is not configured")
    model = os.getenv("GEMINI_IMAGE_MODEL", "gemini-3.1-flash-image").strip()
    payload = {
        "model": model,
        "input": [{"type": "text", "text": prompt}],
        "response_format": {
            "type": "image",
            "mime_type": "image/jpeg",
            "aspect_ratio": aspect_ratio,
            "image_size": os.getenv("GEMINI_IMAGE_SIZE", "1K"),
        },
    }
    with httpx.Client(timeout=httpx.Timeout(300.0, connect=20.0)) as client:
        response = client.post(
            "https://generativelanguage.googleapis.com/v1beta/interactions",
            headers={"x-goog-api-key": api_key, "Content-Type": "application/json"},
            json=payload,
        )
        if not response.is_success:
            try:
                detail = response.json().get("error", {}).get("message", response.text)
            except Exception:
                detail = response.text
            raise ValueError(f"Nano Banana request failed ({response.status_code}): {str(detail)[:500]}")
        data = response.json()
    image_b64 = _find_gemini_image_data(data)
    if not image_b64:
        raise ValueError("Nano Banana returned no image data")
    return image_b64, model

def generate_openai_image(prompt: str) -> tuple[str, str]:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise ValueError("OPENAI_API_KEY is not configured")
    model = os.getenv("OPENAI_IMAGE_MODEL", "gpt-image-1").strip()
    response = OpenAI(api_key=api_key, timeout=300).images.generate(
        model=model,
        prompt=prompt,
        size="1024x1536",
        quality=os.getenv("OPENAI_IMAGE_QUALITY", "medium").strip(),
        output_format="jpeg",
        n=1,
    )
    item = response.data[0] if response.data else None
    image_b64 = getattr(item, "b64_json", "") if item else ""
    if not image_b64 and item and getattr(item, "url", ""):
        image_response = httpx.get(item.url, timeout=120, follow_redirects=True)
        image_response.raise_for_status()
        image_b64 = base64.b64encode(image_response.content).decode("ascii")
    if not image_b64:
        raise ValueError("OpenAI returned no image data")
    return image_b64, model

def _is_quota_error(error: Exception):
    message = str(error).lower()
    return any(marker in message for marker in ("quota", "429", "resource_exhausted", "rate limit", "rate_limit", "exceeded"))

@app.post("/api/generate-image")
def generate_image(req: ImageRequest):
    if not os.getenv("GEMINI_API_KEY") and not os.getenv("OPENAI_API_KEY"):
        raise HTTPException(400, "لم يتم إعداد أي مزود لتوليد الصور")
    try:
        base_prompt = visual_prompt(req)
        image_provider = "openai"
        fallback_used = False
        fallback_reason = ""
        if os.getenv("GEMINI_API_KEY"):
            try:
                image_b64, image_model = generate_nano_banana_image(base_prompt)
                image_provider = "google_nano_banana_2"
            except Exception as gemini_error:
                if not os.getenv("OPENAI_API_KEY"):
                    raise
                fallback_used = True
                fallback_reason = "gemini_quota" if _is_quota_error(gemini_error) else "gemini_error"
                image_b64, image_model = generate_openai_image(base_prompt)
        else:
            image_b64, image_model = generate_openai_image(base_prompt)
        if os.getenv("OPENAI_API_KEY"):
            try:
                review = review_artwork(OpenAI(api_key=os.getenv("OPENAI_API_KEY")), req, image_b64)
            except Exception as review_error:
                review = {"semantic_match":None,"score":None,"issues":[],"retry_direction":"","summary_ar":"تعذر تنفيذ المراجعة البصرية، وتم الاحتفاظ بالصورة المولدة.","review_error":str(review_error)[:300]}
        else:
            review = {"semantic_match":None,"score":None,"issues":[],"retry_direction":"","summary_ar":"لم تُنفذ المراجعة البصرية لأن OPENAI_API_KEY غير مهيأ."}
        return {"b64_json":image_b64,"mime_type":"image/jpeg","slide_number":req.slide_number,"visual_style":req.visual_style,"font":"Cairo","overlay_required":True,"hashtags_in_image":False,"artwork_version":"single-image-provider-fallback-v10","generation_attempts":1,"variant_index":req.variant_index,"image_provider":image_provider,"image_model":image_model,"fallback_used":fallback_used,"fallback_reason":fallback_reason,"semantic_review":review}
    except Exception as e: raise HTTPException(500, f"Image generation failed: {e}")
